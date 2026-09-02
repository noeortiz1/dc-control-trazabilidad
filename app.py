# -*- coding: utf-8 -*-
import os
import json
import base64
import zipfile
import urllib.request
import smtplib
from email.mime.text import MIMEText
from email.mime.multipart import MIMEMultipart
import psycopg2
import psycopg2.pool
import psycopg2.extras
import re
import pandas as pd
import plotly.express as px
from datetime import datetime, date, timedelta
import streamlit as st
import io
import shutil
from io import BytesIO
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
import shutil

# ==========================================
# CONFIGURACIÓN DE PÁGINA Y ESTILO CORPORATIVO
# ==========================================
st.set_page_config(
    page_title="Control de Cotizaciones - DC Control",
    layout="wide",
    page_icon="🏗️"
)

# Estilos CSS Limpios, Ejecutivos y 100% Libres de Monday/Control de Cotizaciones/etc.
st.markdown("""
<style>
    /* Cabecera Principal */
    .system-header {
        background-color: #111827;
        color: white;
        padding: 20px;
        border-radius: 8px;
        margin-bottom: 25px;
        border-left: 8px solid #00C875;
    }
    .system-title {
        font-size: 26px;
        font-weight: 700;
        margin: 0;
    }
    .system-subtitle {
        color: #9ca3af;
        font-size: 13px;
        margin-top: 5px;
    }
    
    /* Tarjetas de Estados / Badges Nativos */
    .badge-status {
        padding: 5px 10px;
        border-radius: 4px;
        font-weight: bold;
        color: white;
        display: inline-block;
        font-size: 12px;
        text-align: center;
    }
    .bg-proceso { background-color: #FDAB3D; }
    .bg-ganado { background-color: #00C875; }
    .bg-perdido { background-color: #E2445C; }
    .bg-cancelado { background-color: #797E93; }
    .bg-revision { background-color: #0085FF; }
</style>
""", unsafe_allow_html=True)

# ==========================================
# CONSTANTES Y CONFIGURACIONES
# ==========================================
DB_FILE = "pipeline_cotizaciones.db"
UPLOAD_DIR = "uploads"

# Asegurar que exista el directorio de cargas
if not os.path.exists(UPLOAD_DIR):
    os.makedirs(UPLOAD_DIR)

ESTADOS_MEXICO = {
    "CDMX": "Líder Regional - Sur",
    "Estado de México": "Líder Regional - Sur",
    "Querétaro": "Líder Regional - Sur",
    "Guanajuato": "Líder Regional - Sur",
    "Jalisco": "Líder Regional - Sur",
    "Michoacán": "Líder Regional - Sur",
    "Puebla": "Líder Regional - Sur",
    "Veracruz": "Líder Regional - Sur",
    "Hidalgo": "Líder Regional - Sur",
    "Morelos": "Líder Regional - Sur",
    "Guerrero": "Líder Regional - Sur",
    "Oaxaca": "Líder Regional - Sur",
    "Chiapas": "Líder Regional - Sur",
    "Tabasco": "Líder Regional - Sur",
    "Campeche": "Líder Regional - Sur",
    "Yucatán": "Líder Regional - Sur",
    "Quintana Roo": "Líder Regional - Sur",
    "Tlaxcala": "Líder Regional - Sur",
    "Colima": "Líder Regional - Sur",
    "Nayarit": "Líder Regional - Sur",
    "Nuevo León": "Líder Regional - Norte",
    "Chihuahua": "Líder Regional - Norte",
    "Coahuila": "Líder Regional - Norte",
    "Sonora": "Líder Regional - Norte",
    "Baja California": "Líder Regional - Norte",
    "Baja California Sur": "Líder Regional - Norte",
    "San Luis Potosí": "Líder Regional - Norte",
    "Aguascalientes": "Líder Regional - Norte",
    "Durango": "Líder Regional - Norte",
    "Sinaloa": "Líder Regional - Norte",
    "Zacatecas": "Líder Regional - Norte",
    "Tamaulipas": "Líder Regional - Norte"
}

# ==========================================
# GESTIÓN DE BASE DE DATOS POSTGRESQL (SUPABASE)
# ==========================================
DB_URI_PRIMARY = "postgresql://postgres.vtakzlcbjizdfbppgrqg:vCQhcfq72BtFbbtx@aws-0-us-east-1.pooler.supabase.com:6543/postgres"
DB_URI_SECONDARY = "postgresql://postgres.vtakzlcbjizdfbppgrqg:vCQhcfq72BtFbbtx@aws-0-us-east-1.pooler.supabase.com:6543/postgres"

class PostgreSQLCursorWrapper:
    def __init__(self, cursor):
        self._cursor = cursor
        
    def execute(self, sql, params=None):
        # Intercept SQLite-specific PRAGMA table_info and redirect to PostgreSQL information_schema
        if "PRAGMA table_info" in sql:
            match = re.search(r"PRAGMA table_info\((.*?)\)", sql)
            if match:
                table_name = match.group(1).replace("'", "").replace('"', '').strip()
                sql_pg = f"SELECT 0 as cid, column_name as name FROM information_schema.columns WHERE table_name = '{table_name}'"
                self._cursor.execute(sql_pg)
                return self
        
        # Replace AUTOINCREMENT with SERIAL for table creation
        if "PRIMARY KEY AUTOINCREMENT" in sql:
            sql = sql.replace("INTEGER PRIMARY KEY AUTOINCREMENT", "SERIAL PRIMARY KEY")
            
        # Replace SQLite ? placeholder with PostgreSQL %s
        sql_pg = sql.replace('?', '%s')
        
        # Ensure params is a tuple/list for psycopg2
        if params is not None:
            if not isinstance(params, (tuple, list)):
                params = (params,)
        
        self._cursor.execute(sql_pg, params)
        return self
        
    def fetchone(self):
        return self._cursor.fetchone()
        
    def fetchall(self):
        return self._cursor.fetchall()
        
    def executemany(self, sql, seq_of_params):
        sql_pg = sql.replace('?', '%s')
        self._cursor.executemany(sql_pg, seq_of_params)
        return self
        
    @property
    def rowcount(self):
        return self._cursor.rowcount
        
    def close(self):
        self._cursor.close()

@st.cache_resource
def get_connection_pool():
    try:
        pool = psycopg2.pool.ThreadedConnectionPool(
            minconn=1,
            maxconn=20,
            dsn=DB_URI_PRIMARY,
            connect_timeout=5
        )
        return pool
    except Exception as e:
        try:
            pool = psycopg2.pool.ThreadedConnectionPool(
                minconn=1,
                maxconn=20,
                dsn=DB_URI_SECONDARY,
                connect_timeout=5
            )
            return pool
        except Exception as ex:
            st.error(f"Error de conexión con el Pooler de Supabase: {ex}")
            raise ex

class PostgreSQLConnectionWrapper:
    def __init__(self, conn, pool=None):
        self._conn = conn
        self._pool = pool
    
    def cursor(self, *args, **kwargs):
        cursor = self._conn.cursor(cursor_factory=psycopg2.extras.DictCursor, *args, **kwargs)
        return PostgreSQLCursorWrapper(cursor)
        
    def execute(self, sql, params=None):
        cursor = self._conn.cursor(cursor_factory=psycopg2.extras.DictCursor)
        cursor_wrapper = PostgreSQLCursorWrapper(cursor)
        cursor_wrapper.execute(sql, params)
        return cursor_wrapper
        
    def commit(self):
        self._conn.commit()
        # Invalidate dashboard cache on any write operation
        try:
            get_cached_dashboard_stats.clear()
        except Exception:
            pass
        
    def close(self):
        if self._pool is not None:
            try:
                self._pool.putconn(self._conn)
            except Exception:
                try:
                    self._conn.close()
                except Exception:
                    pass
        else:
            try:
                self._conn.close()
            except Exception:
                pass



def get_system_setting(key, default=""):
    conn = get_db_connection()
    try:
        row = conn.execute("SELECT val FROM system_settings WHERE key = %s", (key,)).fetchone()
        if row:
            return row['val']
    except Exception:
        pass
    finally:
        conn.close()
    return default

def set_system_setting(key, val):
    conn = get_db_connection()
    try:
        cursor = conn.cursor()
        cursor.execute("INSERT INTO system_settings (key, val) VALUES (%s, %s) ON CONFLICT (key) DO UPDATE SET val = EXCLUDED.val", (key, val))
        conn.commit()
    except Exception:
        pass
    finally:
        conn.close()

def save_uploaded_file(uploaded_file, project_id, step_name):
    # 1. Get raw bytes
    file_bytes = uploaded_file.getvalue()
    
    # 2. Save locally as fallback/local copy
    file_path = os.path.join(UPLOAD_DIR, f"{project_id}_{step_name}_{datetime.now().strftime('%H%M%S')}_{uploaded_file.name}")
    try:
        with open(file_path, "wb") as f_out:
            f_out.write(file_bytes)
    except Exception:
        pass
        
    # 3. Save to database including bytes
    conn = get_db_connection()
    try:
        cursor = conn.cursor()
        cursor.execute('''
            INSERT INTO uploads (project_id, step_name, filename, file_path, uploaded_by, uploaded_at, file_data)
            VALUES (%s, %s, %s, %s, %s, %s, %s)
        ''', (project_id, step_name, uploaded_file.name, file_path, st.session_state.user_name, datetime.now().strftime("%Y-%m-%d %H:%M:%S"), psycopg2.Binary(file_bytes)))
        conn.commit()
    except Exception as e:
        st.error(f"Error al guardar archivo en base de datos: {e}")
    finally:
        conn.close()
    return file_path

def dispatch_step_completion_notifications(project_id, completed_step_num):
    # Check if notifications are enabled
    if get_system_setting("notifications_enabled", "0") != "1":
        return
        
    # Get project details
    conn = get_db_connection()
    p = conn.execute("SELECT * FROM projects WHERE id = %s", (project_id,)).fetchone()
    conn.close()
    if not p:
        return
        
    next_step_num = completed_step_num + 1
    if next_step_num > 7:
        return # Flow completed
        
    # Define next step name, description and assigned role/name
    steps_meta = {
        1: {
            "name": "Paso 1: Levantamiento Técnico",
            "desc": "Cargar la evidencia y datos técnicos del levantamiento de la obra.",
            "assignee": p['assigned_ventas']
        },
        2: {
            "name": "Paso 2: Reunión de Seguimiento y Minuta de Trabajo",
            "desc": "Realizar la reunión comercial-técnica y subir la minuta de trabajo firmada por Ventas y Líder.",
            "assignee": f"Agente Ventas ({p['assigned_ventas']}) y Líder Regional ({p['assigned_lider']})"
        },
        3: {
            "name": "Paso 3: Catálogo de Conceptos Técnico",
            "desc": "Elaborar y subir el catálogo de conceptos técnicos de ingeniería.",
            "assignee": p['assigned_lider']
        },
        4: {
            "name": "Paso 4: Elaboración de Cotización de Precios",
            "desc": "Formular los precios unitarios, márgenes de utilidad y cargar la cotización final.",
            "assignee": p['assigned_costos']
        },
        5: {
            "name": "Paso 5: Revisión de Cotización y Aprobación de Costos",
            "desc": "Revisión a detalle de costos, alcance y margen comercial para su firma autorizada.",
            "assignee": "Dirección General / Comercial / Proyectos"
        },
        6: {
            "name": "Paso 6: Entrega Comercial al Cliente",
            "desc": "Entregar formalmente la propuesta al cliente final y registrar observaciones y monto entregado.",
            "assignee": p['assigned_ventas']
        },
        7: {
            "name": "Paso 7: Cierre Comercial de Licitación",
            "desc": "Especifique el resultado comercial definitivo (Ganado / Perdido / Cancelado).",
            "assignee": "Dirección General"
        }
    }
    
    meta = steps_meta.get(next_step_num)
    if not meta:
        return
        
    # Find email addresses for next step's assignees
    emails = []
    conn = get_db_connection()
    try:
        assignee_name = meta['assignee']
        rows = conn.execute("SELECT email, full_name FROM users WHERE full_name = %s OR role = %s", (assignee_name, assignee_name)).fetchall()
        for r in rows:
            if r['email'] and "@" in r['email']:
                emails.append((r['email'], r['full_name']))
                
        # If no specific email is found and it's step 5, notify all Directors
        if next_step_num == 5:
            directors = conn.execute("SELECT email, full_name FROM users WHERE role LIKE '%Director%'").fetchall()
            for d in directors:
                if d['email'] and "@" in d['email']:
                    emails.append((d['email'], d['full_name']))
    except Exception:
        pass
    finally:
        conn.close()
        
    # --- 1. DISPATCH EMAIL NOTIFICATIONS ---
    smtp_host = get_system_setting("smtp_host")
    smtp_port = get_system_setting("smtp_port")
    smtp_user = get_system_setting("smtp_user")
    smtp_pass = get_system_setting("smtp_pass")
    smtp_sender = get_system_setting("smtp_sender", "DC Control Notificaciones")
    
    if smtp_host and smtp_port and smtp_user and smtp_pass and emails:
        for email, f_name in set(emails): # Deduplicate emails
            try:
                msg = MIMEMultipart()
                msg['From'] = f"{smtp_sender} <{smtp_user}>"
                msg['To'] = email
                msg['Subject'] = f"🏗️ DC Control - Tarea Asignada: {project_id} - {p['name']}"
                
                body = f"""<html>
<body style="font-family: Arial, sans-serif; color: #333333; line-height: 1.6;">
    <div style="background-color: #111827; color: white; padding: 20px; border-radius: 6px 6px 0 0; border-left: 6px solid #00C875;">
        <h2 style="margin: 0; font-size: 20px;">🏗️ Control de Cotizaciones - DC Control</h2>
    </div>
    <div style="padding: 20px; border: 1px solid #e5e7eb; border-top: none; border-radius: 0 0 6px 6px;">
        <p>Hola <strong>{f_name}</strong>,</p>
        <p>Te informamos que se ha completado de manera exitosa el paso anterior en el proyecto <strong>{project_id} - {p['name']}</strong> para el cliente <strong>{p['client']}</strong>.</p>
        
        <p style="background-color: #f3f4f6; padding: 15px; border-radius: 4px; border-left: 4px solid #0085FF;">
            💼 <strong>Siguiente Acción Requerida:</strong><br>
            <span style="font-size: 16px; font-weight: bold; color: #111827;">{meta['name']}</span><br>
            <span style="color: #4b5563;">{meta['desc']}</span>
        </p>
        
        <p><strong>Responsable Asignado:</strong> {meta['assignee']}</p>
        <p><strong>Fecha Límite Compromiso:</strong> {p['target_date']}</p>
        
        <hr style="border: none; border-top: 1px solid #e5e7eb; margin: 20px 0;">
        <p style="font-size: 13px; color: #6b7280;">Por favor, ingresa a la aplicación de escritorio de <strong>DC Control</strong> para continuar con el flujo secuencial y registrar la información técnica correspondiente.</p>
        <p style="text-align: center; margin-top: 25px;">
            <span style="background-color: #00C875; color: white; padding: 10px 20px; border-radius: 4px; text-decoration: none; font-weight: bold; display: inline-block;">DC Control S.A. de C.V.</span>
        </p>
    </div>
</body>
</html>"""
                msg.attach(MIMEText(body, 'html'))
                
                server = smtplib.SMTP(smtp_host, int(smtp_port))
                server.starttls()
                server.login(smtp_user, smtp_pass)
                server.sendmail(smtp_user, email, msg.as_string())
                server.quit()
            except Exception as e:
                pass
                
    # --- 2. DISPATCH TEAMS NOTIFICATION (WEBHOOK) ---
    teams_url = get_system_setting("teams_webhook_url")
    if teams_url and teams_url.startswith("http"):
        try:
            card_payload = {
                "@type": "MessageCard",
                "@context": "http://schema.org/extensions",
                "themeColor": "00C875",
                "summary": f"DC Control - Tarea Asignada {project_id}",
                "sections": [{
                    "activityTitle": f"🏗️ DC Control - Siguiente Paso Habilitado",
                    "activitySubtitle": f"Proyecto: {project_id} - {p['name']}",
                    "activityImage": "https://img.icons8.com/color/96/000000/crane.png",
                    "facts": [
                        {"name": "Cliente:", "value": p['client']},
                        {"name": "Zona / Región:", "value": f"{p['state']} ({p['zone']})"},
                        {"name": "Prioridad:", "value": p.get('priority', 'Media')},
                        {"name": "Siguiente Tarea:", "value": meta['name']},
                        {"name": "Asignado a:", "value": meta['assignee']},
                        {"name": "Fecha Límite:", "value": p['target_date']}
                    ],
                    "markdown": True
                }]
            }
            req = urllib.request.Request(
                teams_url,
                data=json.dumps(card_payload).encode('utf-8'),
                headers={'Content-Type': 'application/json'}
            )
            with urllib.request.urlopen(req) as response:
                pass
        except Exception:
            pass

def generate_structured_zip_backup():
    conn = get_db_connection()
    try:
        projects = conn.execute("SELECT * FROM projects").fetchall()
        uploads = conn.execute("SELECT * FROM uploads").fetchall()
        audit_log = conn.execute("SELECT * FROM audit_log").fetchall()
    finally:
        conn.close()
        
    zip_buffer = io.BytesIO()
    with zipfile.ZipFile(zip_buffer, "w", zipfile.ZIP_DEFLATED) as zip_file:
        for p in projects:
            p_id = p['id']
            p_name = p['name']
            clean_name = re.sub(r'[\\/*?:"<>|]', "", p_name).strip()
            folder_name = f"{p_id} - {clean_name}"
            
            # 1. Generate Word Dossier
            try:
                report_bytes = generate_docx_report(p_id)
                if report_bytes:
                    zip_file.writestr(f"{folder_name}/Dossier_Ejecutivo_{p_id}.docx", report_bytes)
            except Exception:
                pass
                
            # 2. Metadata text file
            try:
                summary_text = f"""==================================================
RESUMEN GENERAL DE LICITACIÓN COMERCIAL - DC CONTROL
==================================================
ID Proyecto: {p_id}
Obra / Proyecto: {p['name']}
Cliente: {p['client']}
Estado / Región: {p['state']} ({p['zone']})
Prioridad: {p.get('priority', 'Media')}
Monto Final Cotizado: ${p['final_amount']:,.2f}
Estatus Comercial: {p['status']}
Paso Actual: {p['current_stage']} de 7

Líder Regional Asignado: {p['assigned_lider']}
Analista de Costos Asignado: {p['assigned_costos']}
Agente de Ventas Asignado: {p['assigned_ventas']}
Fecha de Registro: {p['created_at']}
Fecha Compromiso de Entrega: {p['target_date']}
Motivo de Cierre / Pérdida: {p['lose_reason'] or 'N/A'}
Margen / Desfase de Precio (%): {p['lose_percentage_gap']}%

ESTATUS DE COMPUERTAS SECUENCIALES:
- Paso 1 (Levantamiento): {'Completado' if p['step1_completed'] == 1 else 'Pendiente'}
- Paso 2 (Minuta): {'Completado' if p['step2_completed'] == 1 else 'Pendiente'}
- Paso 3 (Catálogo): {'Completado' if p['step3_completed'] == 1 else 'Pendiente'}
- Paso 4 (Cotización): {'Completado' if p['step4_completed'] == 1 else 'Pendiente'}
- Paso 5 (Revisión Dirección): {'Completado' if p['step5_completed'] == 1 else 'Pendiente'}
- Paso 6 (Entrega Cliente): {'Completado' if p['step6_completed'] == 1 else 'Pendiente'}
=================================================="""
                zip_file.writestr(f"{folder_name}/Resumen_Licitacion_{p_id}.txt", summary_text.encode('utf-8'))
            except Exception:
                pass
                
            # 3. Audit log text file
            try:
                p_logs = [log for log in audit_log if log['project_id'] == p_id]
                log_lines = []
                log_lines.append("==================================================")
                log_lines.append(f"HISTORIAL DE AUDITORÍA Y TRAZABILIDAD - {p_id}")
                log_lines.append("==================================================")
                for l in p_logs:
                    log_lines.append(f"[{l['timestamp']}] User: {l['user_name']} ({l['role']}) - Acción: {l['action']}")
                log_lines.append("==================================================")
                log_text = "\n".join(log_lines)
                zip_file.writestr(f"{folder_name}/Bitacora_Auditoria_{p_id}.txt", log_text.encode('utf-8'))
            except Exception:
                pass
                
            # 4. Binary uploads
            try:
                p_uploads = [up for up in uploads if up['project_id'] == p_id]
                for up in p_uploads:
                    f_content = None
                    if 'file_data' in dict(up) and up['file_data'] is not None:
                        try:
                            f_content = bytes(up['file_data'])
                        except:
                            pass
                    if f_content is None:
                        try:
                            with open(up['file_path'], "rb") as f_in:
                                f_content = f_in.read()
                        except:
                            pass
                            
                    if f_content is not None:
                        clean_fn = re.sub(r'[\\/*?:"<>|]', "_", up['filename']).strip()
                        zip_file.writestr(f"{folder_name}/documentos/{up['step_name']}_{clean_fn}", f_content)
            except Exception:
                pass
                
    zip_buffer.seek(0)
    return zip_buffer.getvalue()

def filter_user_projects(projects_list):
    # If user is not logged in yet or no user_role is set, return list as is
    if 'logged_in' not in st.session_state or not st.session_state.logged_in:
        return projects_list
    role = st.session_state.user_role
    role_clean = str(role).strip().lower()
    full_name = st.session_state.full_name
    
    # If the user is an Admin/Director or any Director role, they see ALL projects
    if role_clean in ["admin/director", "director comercial", "director de proyectos"] or "director" in role_clean:
        return projects_list
        
    filtered = []
    for p in projects_list:
        is_assigned = False
        p_dict = dict(p)
        # Agente de Ventas
        if p_dict.get('assigned_ventas') and (p_dict['assigned_ventas'] == full_name or p_dict['assigned_ventas'] == role):
            is_assigned = True
        # Líder Regional
        if p_dict.get('assigned_lider') and (p_dict['assigned_lider'] == full_name or p_dict['assigned_lider'] == role):
            is_assigned = True
        # Analista de Costos
        if p_dict.get('assigned_costos') and (p_dict['assigned_costos'] == full_name or p_dict['assigned_costos'] == role):
            is_assigned = True
            
        if is_assigned:
            filtered.append(p)
    return filtered

def get_db_connection():

    pool = get_connection_pool()
    conn = pool.getconn()
    return PostgreSQLConnectionWrapper(conn, pool)

def save_file_directly_to_pc(data_bytes, filename):
    import os
    import subprocess
    try:
        user_home = os.path.expanduser('~')
        desktop_path = os.path.join(user_home, 'Desktop')
        if not os.path.exists(desktop_path):
            desktop_path = os.path.join(user_home, 'Escritorio')
        if not os.path.exists(desktop_path):
            desktop_path = os.path.join(user_home, 'Downloads')
        if not os.path.exists(desktop_path):
            desktop_path = user_home
            
        full_dest_path = os.path.join(desktop_path, filename)
        with open(full_dest_path, "wb") as f_out:
            f_out.write(data_bytes)
            
        st.success(f"💾 **¡Guardado con éxito!** Archivo disponible en tu Escritorio: `{full_dest_path}`")
        
        try:
            os.startfile(full_dest_path)
        except:
            try:
                subprocess.Popen(f'explorer /select,"{full_dest_path}"', shell=True)
            except:
                pass
    except Exception as e_direct:
        st.error(f"Error al guardar directamente en PC: {e_direct}")

def init_db(insert_demos=False):
    conn = get_db_connection()
    cursor = conn.cursor()
    
    # 1. Tabla de Proyectos (Esquema de Flujo Secuencial)
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS projects (
            id TEXT PRIMARY KEY,
            name TEXT NOT NULL,
            client TEXT,
            total_amount REAL DEFAULT 0.0,
            final_amount REAL DEFAULT 0.0,
            state TEXT,
            zone TEXT,
            assigned_lider TEXT,
            assigned_costos TEXT,
            assigned_ventas TEXT,
            assigned_ventas_2 TEXT,
            priority TEXT DEFAULT 'Media',
            status TEXT DEFAULT 'En Proceso',
            current_stage INTEGER DEFAULT 1,
            lose_reason TEXT,
            lose_percentage_gap REAL DEFAULT 0.0,
            created_at TEXT,
            target_date TEXT,
            step1_completed INTEGER DEFAULT 0,
            step2_ventas_done INTEGER DEFAULT 0,
            step2_lider_done INTEGER DEFAULT 0,
            step2_completed INTEGER DEFAULT 0,
            step3_completed INTEGER DEFAULT 0,
            step4_completed INTEGER DEFAULT 0,
            step5_completed INTEGER DEFAULT 0,
            step6_completed INTEGER DEFAULT 0
        )
    ''')
    
    # 2. Tabla de Archivos Adjuntos
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS uploads (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id TEXT,
            step_name TEXT,
            filename TEXT,
            file_path TEXT,
            uploaded_by TEXT,
            uploaded_at TEXT
        )
    ''')
    
    # 3. Tabla de Usuarios (with email column included from scratch for new DBs)
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS users (
            username TEXT PRIMARY KEY,
            password TEXT,
            full_name TEXT,
            role TEXT,
            email TEXT
        )
    ''')
    
    # 4. Tabla de Auditoría
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS audit_log (
            id INTEGER PRIMARY KEY AUTOINCREMENT,
            project_id TEXT,
            user_name TEXT,
            role TEXT,
            action TEXT,
            timestamp TEXT
        )
    ''')
    
    # 5. Tabla de Configuración de Sistema (system_settings)
    cursor.execute('''
        CREATE TABLE IF NOT EXISTS system_settings (
            key TEXT PRIMARY KEY,
            val TEXT
        )
    ''')
    
    # --- PROCESO DE MIGRACIÓN AUTÓNOMA ULTRA-ROBUSTA (Auto-Healing Individual) ---
    def get_columns(table_name):
        try:
            cursor.execute(f"SELECT column_name FROM information_schema.columns WHERE table_name = '{table_name}'")
            cols = [row[0] for row in cursor.fetchall()]
            if cols:
                return cols
        except Exception:
            pass
        try:
            cursor.execute(f"PRAGMA table_info({table_name})")
            return [row[1] for row in cursor.fetchall()]
        except Exception:
            return []
            
    proj_cols = get_columns('projects')
    needed_cols = {
        'final_amount': 'REAL DEFAULT 0.0',
        'assigned_ventas': 'TEXT',
        'assigned_ventas_2': 'TEXT',
        'priority': "TEXT DEFAULT 'Media'",
        'lose_reason': 'TEXT',
        'lose_percentage_gap': 'REAL DEFAULT 0.0',
        'target_date': 'TEXT',
        'step1_completed': 'INTEGER DEFAULT 0',
        'step2_ventas_done': 'INTEGER DEFAULT 0',
        'step2_lider_done': 'INTEGER DEFAULT 0',
        'step2_completed': 'INTEGER DEFAULT 0',
        'step3_completed': 'INTEGER DEFAULT 0',
        'step4_completed': 'INTEGER DEFAULT 0',
        'step5_completed': 'INTEGER DEFAULT 0',
        'step6_completed': 'INTEGER DEFAULT 0'
    }
    
    # Migrate projects columns one by one in its own try-catch
    for col_name, col_type in needed_cols.items():
        if col_name not in proj_cols:
            try:
                cursor.execute(f"ALTER TABLE projects ADD COLUMN {col_name} {col_type}")
            except Exception as e:
                pass # Silently proceed to next columns if any error occurs
                
    # Ensure uploads table has uploaded_by column
    upload_cols = get_columns('uploads')
    if 'uploaded_by' not in upload_cols:
        try:
            cursor.execute("ALTER TABLE uploads ADD COLUMN uploaded_by TEXT")
        except Exception:
            pass
            
    # Ensure users has email column
    user_cols = get_columns('users')
    if 'email' not in user_cols:
        try:
            cursor.execute("ALTER TABLE users ADD COLUMN email TEXT")
        except Exception:
            pass
            
    # Ensure uploads table has file_data column (BYTEA in postgres, BLOB in sqlite)
    upload_cols = get_columns('uploads')
    if 'file_data' not in upload_cols:
        try:
            cursor.execute("ALTER TABLE uploads ADD COLUMN file_data BYTEA")
        except Exception:
            try:
                cursor.execute("ALTER TABLE uploads ADD COLUMN file_data BLOB")
            except Exception:
                pass

    # Clean up ALL old demo accounts and previous admin account to make space for the official Noe Ortiz admin
    try:
        cursor.execute("DELETE FROM users WHERE username IN ('admin', 'ventas1', 'lider_sur', 'lider_norte', 'costos_jefe', 'costos_jr1', 'costos_jr2', 'ingeniero')")
    except Exception:
        pass
        
    # Always ensure the admin account is set to the requested Noe Ortiz with the secure credentials (if it doesn't already exist)
    try:
        cursor.execute("SELECT COUNT(*) FROM users WHERE username = 'noe.ortizadm'")
        if cursor.fetchone()[0] == 0:
            cursor.execute("""
                INSERT INTO users (username, password, full_name, role, email)
                VALUES ('noe.ortizadm', 'jaeldiaz251', 'Noe Ortiz (Director General)', 'Admin/Director', 'director@dccontrol.com')
            """)
    except Exception as e:
        pass
        
    # Since the user wants a clean production start, we default insert_demos=False.
    # No demo projects or other users will be inserted, keeping it 100% clean.
    if insert_demos:
        try:
            cursor.execute("SELECT COUNT(*) FROM projects")
            if cursor.fetchone()[0] == 0:
                demo_projects = [
                    ("DCC-202608-N-001", "Ampliación Planta Monterrey", "Aceros de Monterrey S.A.", 1250000.0, 1250000.0, "Nuevo León", "Norte", "Ing. Alejandro Mendoza", "Analista de Costos Jefe", "Ing. Carlos", "En Proceso", 4, None, 0.0, "2026-08-01", "2026-09-15", 1, 1, 1, 1, 1, 0, 0, 0),
                    ("DCC-202608-S-001", "Instalación Eléctrica Querétaro", "Logística del Centro", 450000.0, 450000.0, "Querétaro", "Sur", "Ing. Sofía Romero", "Analista de Costos Junior 1", "Ing. Carlos", "Ganado", 7, None, 0.0, "2026-08-05", "2026-10-10", 1, 1, 1, 1, 1, 1, 1, 1)
                ]
                cursor.executemany("INSERT INTO projects VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)", demo_projects)
        except Exception:
            pass
            
    conn.commit()
    conn.close()

init_db(insert_demos=False)

def log_audit(project_id, user_name, role, action):
    conn = get_db_connection()
    conn.execute('''
        INSERT INTO audit_log (project_id, user_name, role, action, timestamp)
        VALUES (?, ?, ?, ?, ?)
    ''', (project_id, user_name, role, action, datetime.now().strftime("%Y-%m-%d %H:%M:%S")))
    conn.commit()
    conn.close()

def generate_next_project_id(state):
    # Obtener prefijos
    year_month = datetime.now().strftime("%Y%m")
    region_auto = ESTADOS_MEXICO[state]
    zone_code = "S" if "Sur" in region_auto else "N"
    prefix = f"DCC-{year_month}-{zone_code}-"
    
    # Consultar último consecutivo
    conn = get_db_connection()
    row = conn.execute("SELECT id FROM projects WHERE id LIKE ? ORDER BY id DESC LIMIT 1", (prefix + "%",)).fetchone()
    conn.close()
    
    if row:
        last_id = row['id']
        try:
            last_num = int(last_id.split("-")[-1])
            next_num = last_num + 1
        except:
            next_num = 1
    else:
        next_num = 1
        
    return f"{prefix}{next_num:03d}"

# ==========================================
# GENERACIÓN DE DOCUMENTO DOCX (DOSSIER EJECUTIVO)
# ==========================================
def generate_docx_report(project_id):
    conn = get_db_connection()
    p = conn.execute("SELECT * FROM projects WHERE id = ?", (project_id,)).fetchone()
    uploads = conn.execute("SELECT * FROM uploads WHERE project_id = ? ORDER BY uploaded_at ASC", (project_id,)).fetchall()
    logs = conn.execute("SELECT * FROM audit_log WHERE project_id = ? ORDER BY timestamp DESC", (project_id,)).fetchall()
    conn.close()

    if not p:
        return None

    doc = Document()
    
    # Configuración de márgenes
    for section in doc.sections:
        section.top_margin = Inches(1.0)
        section.bottom_margin = Inches(1.0)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # Estilos globales
    style_normal = doc.styles['Normal']
    style_normal.font.name = 'Arial'
    style_normal.font.size = Pt(11)
    style_normal.font.color.rgb = RGBColor(51, 51, 51)

    # Encabezado
    section = doc.sections[0]
    header = section.header
    header.is_linked_to_previous = False
    hp = header.paragraphs[0]
    hp.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    hrun = hp.add_run("DC Control S.A. de C.V. | Reporte Ejecutivo de Trazabilidad")
    hrun.font.size = Pt(8.5)
    hrun.font.color.rgb = RGBColor(121, 126, 147)

    # Pie de página con números
    footer = section.footer
    footer.is_linked_to_previous = False
    fp = footer.paragraphs[0]
    fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
    frun = fp.add_run("DC Control S.A. de C.V. — Control de Cotizaciones  •  Página ")
    frun.font.size = Pt(9)
    frun.font.color.rgb = RGBColor(121, 126, 147)
    
    fldSimple = OxmlElement('w:fldSimple')
    fldSimple.set(qn('w:instr'), 'PAGE')
    fp._p.append(fldSimple)

    # Título Principal
    title_p = doc.add_paragraph()
    title_p.paragraph_format.space_before = Pt(12)
    title_p.paragraph_format.space_after = Pt(6)
    title_run = title_p.add_run("REPORTE EJECUTIVO DE COTIZACIÓN")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor(17, 24, 39)

    # Línea de separación verde
    bar_p = doc.add_paragraph()
    bar_p.paragraph_format.space_before = Pt(0)
    bar_p.paragraph_format.space_after = Pt(18)
    bar_run = bar_p.add_run("―" * 50)
    bar_run.bold = True
    bar_run.font.color.rgb = RGBColor(0, 200, 117)

    # Sección 1: Resumen General
    h1 = doc.add_paragraph()
    h1_run = h1.add_run("1. Resumen General de la Licitación")
    h1_run.bold = True
    h1_run.font.size = Pt(13)
    h1_run.font.color.rgb = RGBColor(31, 41, 55)
    h1.paragraph_format.space_after = Pt(10)

    # Tabla de datos generales
    ventas_info = str(p['assigned_ventas'])
        
    details = [
        ("Código del Proyecto", str(p['id'])),
        ("Obra / Proyecto", str(p['name'])),
        ("Cliente", str(p['client'])),
        ("Estado de la República", f"{p['state']} ({p['zone']})"),
        ("Agente de Ventas", ventas_info),
        ("Prioridad", str(p.get('priority', 'Media') if p.get('priority') else 'Media')),
        ("Monto Final Cotizado", f"${p['final_amount']:,.2f}"),
        ("Estatus Comercial", str(p['status']))
    ]
    
    table = doc.add_table(rows=len(details), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    for i, (label, val) in enumerate(details):
        row = table.rows[i]
        cell_lbl = row.cells[0]
        cell_lbl.width = Inches(2.2)
        p_lbl = cell_lbl.paragraphs[0]
        p_lbl.add_run(label).bold = True
        
        cell_val = row.cells[1]
        cell_val.width = Inches(4.3)
        p_val = cell_val.paragraphs[0]
        p_val.add_run(val)

        for cell in (cell_lbl, cell_val):
            cell.paragraphs[0].paragraph_format.space_after = Pt(4)
            cell.paragraphs[0].paragraph_format.space_before = Pt(4)
            if cell == cell_lbl:
                shd = OxmlElement('w:shd')
                shd.set(qn('w:fill'), 'F3F4F6')
                shd.set(qn('w:val'), 'clear')
                cell._tc.get_or_add_tcPr().append(shd)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Sección 2: Trazabilidad de Pasos
    h2 = doc.add_paragraph()
    h2_run = h2.add_run("2. Estado de Pasos y Validación")
    h2_run.bold = True
    h2_run.font.size = Pt(13)
    h2_run.font.color.rgb = RGBColor(31, 41, 55)
    h2.paragraph_format.space_after = Pt(10)

    steps = [
        ("Paso 1: Levantamiento Técnico", str(p['assigned_ventas']), p['step1_completed'] == 1),
        ("Paso 2: Minuta de Trabajo", f"{p['assigned_ventas']} & {p['assigned_lider']}", p['step2_completed'] == 1),
        ("Paso 3: Catálogo de Conceptos", str(p['assigned_lider']), p['step3_completed'] == 1),
        ("Paso 4: Elaboración de Cotización", str(p['assigned_costos']), p['step4_completed'] == 1),
        ("Paso 5: Revisión de Dirección", "Dirección General", p['step5_completed'] == 1),
        ("Paso 6: Entrega al Cliente", str(p['assigned_ventas']), p['step6_completed'] == 1),
        ("Paso 7: Cierre Comercial", "Dirección General", p['status'] != 'En Proceso')
    ]

    for step_title, resp, is_done in steps:
        sp = doc.add_paragraph(style='List Bullet')
        sp.paragraph_format.space_after = Pt(4)
        run_title = sp.add_run(f"{step_title}: ")
        run_title.bold = True
        
        sp.add_run(f"({resp}) ― ")
        
        status_text = "✔️ COMPLETADO" if is_done else "⏳ PENDIENTE"
        status_run = sp.add_run(status_text)
        status_run.bold = True
        if is_done:
            status_run.font.color.rgb = RGBColor(0, 200, 117)
        else:
            status_run.font.color.rgb = RGBColor(226, 68, 92)

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Sección 3: Dossier de Archivos
    h3 = doc.add_paragraph()
    h3_run = h3.add_run("3. Expediente Digital de Archivos")
    h3_run.bold = True
    h3_run.font.size = Pt(13)
    h3_run.font.color.rgb = RGBColor(31, 41, 55)
    h3.paragraph_format.space_after = Pt(10)

    if not uploads:
        doc.add_paragraph("No se han cargado documentos para esta cotización aún.")
    else:
        table_u = doc.add_table(rows=len(uploads) + 1, cols=4)
        table_u.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_u.autofit = False

        headers_u = ["Paso / Etapa", "Nombre del Archivo", "Cargado Por", "Fecha de Carga"]
        widths_u = [Inches(1.8), Inches(2.2), Inches(1.3), Inches(1.2)]

        hdr_row = table_u.rows[0]
        for col_idx, text in enumerate(headers_u):
            cell = hdr_row.cells[col_idx]
            cell.width = widths_u[col_idx]
            p_hdr = cell.paragraphs[0]
            p_hdr.paragraph_format.space_before = Pt(4)
            p_hdr.paragraph_format.space_after = Pt(4)
            p_hdr.add_run(text).bold = True
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '111827')
            shd.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shd)
            p_hdr.runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for row_idx, f_data in enumerate(uploads):
            row = table_u.rows[row_idx + 1]
            step_friendly = f_data['step_name'].replace("step1_", "Paso 1: ").replace("step2_", "Paso 2: ").replace("step3_", "Paso 3: ").replace("step4_", "Paso 4: ").replace("step6_", "Paso 6: ")
            row_data = [
                step_friendly.capitalize(),
                f_data['filename'],
                f_data['uploaded_by'],
                f_data['uploaded_at']
            ]
            for col_idx, text in enumerate(row_data):
                cell = row.cells[col_idx]
                cell.width = widths_u[col_idx]
                p_cell = cell.paragraphs[0]
                p_cell.paragraph_format.space_before = Pt(4)
                p_cell.paragraph_format.space_after = Pt(4)
                p_cell.add_run(str(text))

    doc.add_paragraph().paragraph_format.space_after = Pt(12)

    # Sección 4: Historial de Auditoría
    h4 = doc.add_paragraph()
    h4_run = h4.add_run("4. Historial de Cambios (Audit Trail)")
    h4_run.bold = True
    h4_run.font.size = Pt(13)
    h4_run.font.color.rgb = RGBColor(31, 41, 55)
    h4.paragraph_format.space_after = Pt(10)

    if not logs:
        doc.add_paragraph("No se cuenta con registros de auditoría para esta licitación.")
    else:
        table_l = doc.add_table(rows=len(logs) + 1, cols=4)
        table_l.alignment = WD_TABLE_ALIGNMENT.CENTER
        table_l.autofit = False

        headers_l = ["Fecha y Hora", "Colaborador", "Puesto / Rol", "Acción Realizada"]
        widths_l = [Inches(1.2), Inches(1.3), Inches(1.3), Inches(2.7)]

        hdr_row = table_l.rows[0]
        for col_idx, text in enumerate(headers_l):
            cell = hdr_row.cells[col_idx]
            cell.width = widths_l[col_idx]
            p_hdr = cell.paragraphs[0]
            p_hdr.paragraph_format.space_before = Pt(4)
            p_hdr.paragraph_format.space_after = Pt(4)
            p_hdr.add_run(text).bold = True
            shd = OxmlElement('w:shd')
            shd.set(qn('w:fill'), '111827')
            shd.set(qn('w:val'), 'clear')
            cell._tc.get_or_add_tcPr().append(shd)
            p_hdr.runs[0].font.color.rgb = RGBColor(255, 255, 255)

        for row_idx, log in enumerate(logs):
            row = table_l.rows[row_idx + 1]
            row_data = [
                log['timestamp'],
                log['user_name'],
                log['role'],
                log['action']
            ]
            for col_idx, text in enumerate(row_data):
                cell = row.cells[col_idx]
                cell.width = widths_l[col_idx]
                p_cell = cell.paragraphs[0]
                p_cell.paragraph_format.space_before = Pt(4)
                p_cell.paragraph_format.space_after = Pt(4)
                p_cell.add_run(str(text))

    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio.getvalue()


# ==========================================
# INICIO DE SESIÓN
# ==========================================
if 'logged_in' not in st.session_state:
    st.session_state.logged_in = False
    st.session_state.user_name = ""
    st.session_state.user_role = ""
    st.session_state.full_name = ""

# Reset de seguridad para el pin de edición de perfil antes de instanciar el widget
if st.session_state.get('clear_profile_pin', False):
    st.session_state["sec_key_prof"] = ""
    st.session_state['clear_profile_pin'] = False

def logout():
    st.session_state.logged_in = False
    st.session_state.user_name = ""
    st.session_state.user_role = ""
    st.session_state.full_name = ""
    st.rerun()

if not st.session_state.logged_in:
    st.markdown("<br><br>", unsafe_allow_html=True)
    col_l1, col_l2, col_l3 = st.columns([1, 1.2, 1])
    with col_l2:
        with st.container(border=True):
            # Cargar logo si existe en el directorio de trabajo
            if os.path.exists("logo.png"):
                # Squeezed columns inside container to make logo centered and beautifully resized (not oversized)
                col_logo_l, col_logo_c, col_logo_r = st.columns([1, 1.5, 1])
                with col_logo_c:
                    st.image("logo.png", use_container_width=True)
            else:
                st.markdown("<h1 style='text-align: center; color: #111827; margin-bottom: 20px;'>🏗️ DC Control</h1>", unsafe_allow_html=True)
                
            st.markdown("<h4 style='text-align: center; color: #4b5563; margin-top:0;'>Control de Cotizaciones</h4>", unsafe_allow_html=True)
            st.write("---")
            username_input = st.text_input("Usuario")
            password_input = st.text_input("Contraseña", type="password")
            btn_login = st.button("Ingresar al Sistema 🚀", use_container_width=True)
            
            if btn_login:
                conn = get_db_connection()
                user = conn.execute("SELECT * FROM users WHERE username = ? AND password = ?", (username_input, password_input)).fetchone()
                conn.close()
                
                if user:
                    st.session_state.logged_in = True
                    st.session_state.user_name = user['username']
                    st.session_state.user_role = user['role']
                    st.session_state.full_name = user['full_name']
                    st.success("Acceso autorizado con éxito")
                    st.rerun()
                else:
                    st.error("Credenciales incorrectas")
    st.stop()

# ==========================================
# BARRA LATERAL (SIDEBAR)
# ==========================================
if os.path.exists("logo.png"):
    st.sidebar.image("logo.png", use_container_width=True)
    st.sidebar.markdown("---")
else:
    st.sidebar.markdown("### 🏗️ DC Control")

st.sidebar.markdown(f"#### 👤 {st.session_state.full_name}")
st.sidebar.markdown(f"**Puesto:** {st.session_state.user_role}")

if st.sidebar.button("🚪 Cerrar Sesión", use_container_width=True):
    logout()

st.sidebar.markdown("---")
if st.sidebar.button("🔄 Sincronizar / Actualizar Datos", use_container_width=True):
    st.rerun()
st.sidebar.caption("💡 Haz clic para actualizar la pantalla y ver las cargas o cambios que haga tu equipo en tiempo real sin tener que cerrar sesión.")

# ==========================================
# CABECERA EJECUTIVA EN PANTALLA
# ==========================================
col_header_logo, col_header_text = st.columns([1, 6])
with col_header_logo:
    if os.path.exists("logo.png"):
        st.image("logo.png", width=100)
    else:
        st.markdown("<h1 style='margin:0; text-align:center;'>🏗️</h1>", unsafe_allow_html=True)

with col_header_text:
    st.markdown(f"""
    <div style='background-color: #111827; color: white; padding: 15px 20px; border-radius: 8px; border-left: 8px solid #00C875;'>
        <h2 style='margin:0; font-size:22px;'>Control de Cotizaciones</h2>
        <p style='margin:5px 0 0 0; color: #9ca3af; font-size:12px;'>DC Control S.A. de C.V. • Rol activo: {st.session_state.user_role}</p>
    </div>
    """, unsafe_allow_html=True)

# Definición de pestañas
role = st.session_state.user_role
role_clean = str(role).strip().lower()
if role_clean == "admin/director":
    tabs_config = ["📊 Dashboard", "📋 Tablero de Proyectos", "✔️ Compuertas Técnicas", "🗺️ Kanban Visual", "👥 Usuarios y Seguridad", "⚙️ Consola de Control", "📜 Bitácora Auditoría"]
else:
    tabs_config = ["📊 Dashboard", "📋 Tablero de Proyectos", "✔️ Compuertas Técnicas", "🗺️ Kanban Visual"]

tabs = st.tabs(tabs_config)
tab_dict = {name: tab_obj for name, tab_obj in zip(tabs_config, tabs)}

@st.cache_data(ttl=30)
def get_cached_dashboard_stats():
    conn = get_db_connection()
    try:
        total_p = conn.execute("SELECT COUNT(*) FROM projects").fetchone()[0]
        monto_total = conn.execute("SELECT SUM(final_amount) FROM projects").fetchone()[0] or 0.0
        ganados_n = conn.execute("SELECT COUNT(*) FROM projects WHERE status = 'Ganado'").fetchone()[0]
        perdidos_n = conn.execute("SELECT COUNT(*) FROM projects WHERE status = 'Perdido'").fetchone()[0]
        projs_all = [dict(row) for row in conn.execute("SELECT * FROM projects").fetchall()]
        alertas_p = conn.execute("SELECT COUNT(*) FROM projects WHERE status = 'En Proceso'").fetchone()[0]
        return {
            "total_p": total_p,
            "monto_total": monto_total,
            "ganados_n": ganados_n,
            "perdidos_n": perdidos_n,
            "projs_all": projs_all,
            "alertas_p": alertas_p
        }
    finally:
        conn.close()

# ==========================================
# MÓDULO 1: DASHBOARD EJECUTIVO
# ==========================================
if "📊 Dashboard" in tab_dict:
    with tab_dict["📊 Dashboard"]:
        st.subheader("📊 Panel de Control de Cotizaciones")
        
        # Conexión DB y extracción rápida (Cached & Optimized)
        stats = get_cached_dashboard_stats()
        total_p = stats["total_p"]
        monto_total = stats["monto_total"]
        ganados_n = stats["ganados_n"]
        perdidos_n = stats["perdidos_n"]
        projs_all = stats["projs_all"]
        alertas_p = stats["alertas_p"]
        
        # Tarjetas KPI elegantes nativas
        col_k1, col_k2, col_k3, col_k4 = st.columns(4)
        with col_k1:
            st.metric("Licitaciones Registradas", total_p)
        with col_k2:
            st.metric("Monto Total Cotizado", f"${monto_total:,.2f}")
        with col_k3:
            denom_exito = (ganados_n + perdidos_n)
            tasa_exito = (ganados_n / denom_exito * 100) if denom_exito > 0 else 0.0
            st.metric("Efectividad Comercial", f"{tasa_exito:.1f}%")
        with col_k4:
            st.metric("Cotizaciones en Curso", alertas_p)
            
        st.markdown("---")
        
        df_p = pd.DataFrame([dict(p) for p in projs_all]) if projs_all else pd.DataFrame()
        
        # --- SECCIÓN DE ANÁLISIS DE CUELLOS DE BOTELLA PARA REUNIONES DE 5 MINUTOS ---
        st.markdown("### 📋 Resumen")
        st.caption("Estatus de todas las cotizaciones activas en curso. Permite a la Dirección identificar retrasos al instante.")
        
        df_active_bottlenecks = df_p[df_p['status'] == 'En Proceso'].copy() if not df_p.empty else pd.DataFrame()
        
        if df_active_bottlenecks.empty:
            st.success("🎉 ¡No hay proyectos en proceso! Todas las cotizaciones están cerradas o no hay proyectos en curso.")
        else:
            bottlenecks_data = []
            for idx, r_b in df_active_bottlenecks.iterrows():
                # Days left calculation
                days_left = 999
                try:
                    tgt_dt = datetime.strptime(r_b['target_date'], "%Y-%m-%d").date()
                    days_left = (tgt_dt - date.today()).days
                except:
                    pass
                
                # Casing for days remaining
                if days_left < 0:
                    urgencia = f"🔴 Vencido ({abs(days_left)} días)"
                elif days_left <= 7:
                    urgencia = f"🟠 Urgente ({days_left} días)"
                else:
                    urgencia = f"🟢 En Tiempo ({days_left} días)"
                
                # Map step to current stage and responsible
                stg = r_b['current_stage']
                steps_desc_map = {
                    1: "Paso 1: Levantamiento",
                    2: "Paso 2: Minuta Trabajo",
                    3: "Paso 3: Catálogo Conceptos",
                    4: "Paso 4: Cotización Precios",
                    5: "Paso 5: Revisión Dirección",
                    6: "Paso 6: Entrega Cliente",
                    7: "Paso 7: Cierre Comercial"
                }
                step_name_friendly = steps_desc_map.get(stg, f"Paso {stg}")
                
                # Responsible mapping
                if stg == 1:
                    resp_name = r_b.get('assigned_ventas', 'Sin asignar')
                elif stg == 2:
                    ventas_str = r_b.get('assigned_ventas', 'Ventas')
                    lider_str = r_b.get('assigned_lider', 'Líder Regional')
                    resp_name = f"{ventas_str} / {lider_str}"
                elif stg == 3:
                    resp_name = r_b.get('assigned_lider', 'Líder Regional')
                elif stg == 4:
                    resp_name = r_b.get('assigned_costos', 'Analista de Costos')
                elif stg == 5:
                    resp_name = "Dirección General / Directores"
                elif stg == 6:
                    resp_name = r_b.get('assigned_ventas', 'Ventas')
                else:
                    resp_name = "Dirección General"
                    
                ventas_team = r_b.get('assigned_ventas', 'Sin asignar')
                
                bottlenecks_data.append({
                    "ID Licitación": r_b['id'],
                    "Obra / Proyecto": r_b['name'],
                    "Cliente": r_b['client'],
                    "Prioridad": r_b.get('priority', 'Media') if r_b.get('priority') else 'Media',
                    "Paso Atorado": step_name_friendly,
                    "Responsable Actual": resp_name,
                    "Equipo de Ventas": ventas_team,
                    "Fecha Compromiso": r_b['target_date'],
                    "Estatus / Urgencia": urgencia,
                    "days_sort": days_left
                })
                
            df_b_display = pd.DataFrame(bottlenecks_data)
            df_b_display = df_b_display.sort_values(by="days_sort", ascending=True).drop(columns=["days_sort"])
            
            p_options = ["Todos"] + [f"{r['ID Licitación']} - {r['Obra / Proyecto']}" for idx, r in df_b_display.iterrows()]
            b_filter = st.selectbox("📂 Filtrar Resumen por Obra / Proyecto:", p_options, key="bottlenecks_filter")
            if b_filter != "Todos":
                b_filter_id = b_filter.split(" - ")[0]
                df_b_display = df_b_display[df_b_display['ID Licitación'] == b_filter_id]
            
            st.dataframe(
                df_b_display,
                use_container_width=True,
                hide_index=True,
                column_config={
                    "Prioridad": st.column_config.TextColumn(help="Prioridad asignada por Dirección"),
                    "Estatus / Urgencia": st.column_config.TextColumn(help="Semáforo de vencimiento"),
                }
            )
            
        st.markdown("---")
        
        # df_p already defined above
        
        if df_p.empty:
            st.info("No hay datos de cotizaciones registrados de momento. Comience a registrar proyectos para ver el panel de gráficas.")
        else:
            # Row 1 of Charts
            col_g1, col_g2 = st.columns(2)
            
            with col_g1:
                st.markdown("##### 📍 Monto Cotizado por Región")
                fig_reg = px.bar(
                    df_p, 
                    x="zone", 
                    y="final_amount", 
                    color="zone", 
                    labels={"final_amount": "Monto Final ($)", "zone": "Región"},
                    color_discrete_map={"Norte": "#3b82f6", "Sur": "#8b5cf6"},
                    text_auto='.2s'
                )
                fig_reg.update_layout(showlegend=False, height=280, margin=dict(t=10, b=10, l=10, r=10))
                st.plotly_chart(fig_reg, use_container_width=True)
                
            with col_g2:
                st.markdown("##### 📈 Distribución Comercial de Cotizaciones")
                df_p['Estatus Vista'] = df_p['status'].map({
                    'En Proceso': 'En Proceso / Recibidas',
                    'Ganado': 'Ganadas',
                    'Perdido': 'Perdidas',
                    'Cancelado': 'Canceladas / Declinadas'
                }).fillna(df_p['status'])
                
                fig_stat = px.pie(
                    df_p, 
                    names="Estatus Vista", 
                    color="Estatus Vista",
                    color_discrete_map={
                        "En Proceso / Recibidas": "#f59e0b", 
                        "Ganadas": "#10b981", 
                        "Perdidas": "#ef4444", 
                        "Canceladas / Declinadas": "#6b7280"
                    },
                    hole=0.4
                )
                fig_stat.update_layout(height=280, margin=dict(t=10, b=10, l=10, r=10))
                st.plotly_chart(fig_stat, use_container_width=True)
            
            # Row 2 of Charts
            st.markdown("---")
            col_g3, col_g4 = st.columns(2)
            
            with col_g3:
                st.markdown("##### 🏆 Zonas de Mayor Éxito (Licitaciones Ganadas por Estado)")
                df_won = df_p[df_p['status'] == 'Ganado']
                if df_won.empty:
                    st.caption("No se han registrado cotizaciones 'Ganadas' para graficar el éxito por zona.")
                else:
                    df_won_by_state = df_won.groupby('state').size().reset_index(name='Ganadas')
                    fig_success_zones = px.bar(
                        df_won_by_state,
                        x="state",
                        y="Ganadas",
                        labels={"state": "Estado", "Ganadas": "Licitaciones"},
                        color_discrete_sequence=["#10b981"],
                        text_auto=True
                    )
                    fig_success_zones.update_layout(height=280, margin=dict(t=10, b=10, l=10, r=10))
                    st.plotly_chart(fig_success_zones, use_container_width=True)
                    
            with col_g4:
                st.markdown("##### ⚙️ Análisis de Cuello de Botella (Cotizaciones Activas por Paso)")
                df_active = df_p[df_p['status'] == 'En Proceso']
                if df_active.empty:
                    st.caption("No hay cotizaciones activas 'En Proceso' en este momento.")
                else:
                    # Map step names
                    steps_short = {
                        1: "1. Levantamiento",
                        2: "2. Minuta",
                        3: "3. Catálogo",
                        4: "4. Cotización",
                        5: "5. Revisión",
                        6: "6. Entrega",
                        7: "7. Cierre"
                    }
                    df_active['Etapa Nombre'] = df_active['current_stage'].map(steps_short)
                    df_bottleneck = df_active.groupby(['current_stage', 'Etapa Nombre']).size().reset_index(name='Cantidad')
                    df_bottleneck = df_bottleneck.sort_values(by='current_stage')
                    
                    fig_bottleneck = px.bar(
                        df_bottleneck,
                        x="Etapa Nombre",
                        y="Cantidad",
                        labels={"Etapa Nombre": "Paso del Proceso", "Cantidad": "Proyectos Atorados"},
                        color="Cantidad",
                        color_continuous_scale="OrRd",
                        text_auto=True
                    )
                    fig_bottleneck.update_layout(height=280, margin=dict(t=10, b=10, l=10, r=10), showlegend=False, coloraxis_showscale=False)
                    st.plotly_chart(fig_bottleneck, use_container_width=True)
            
            # Row 3: Gap/Desfase and Delivery Warnings
            st.markdown("---")
            col_g5, col_g6 = st.columns(2)
            
            with col_g5:
                st.markdown("##### 📉 Porcentaje de Desfase por Proyecto")
                df_lost = df_p[(df_p['status'] == 'Perdido') & (df_p['lose_percentage_gap'] > 0)]
                if df_lost.empty:
                    st.caption("No hay cotizaciones 'Perdidas' con porcentaje de desfase registrado.")
                else:
                    fig_gap = px.bar(
                        df_lost,
                        x="id",
                        y="lose_percentage_gap",
                        color="zone",
                        labels={"lose_percentage_gap": "Desfase (%)", "id": "Proyecto ID"},
                        color_discrete_map={"Norte": "#3b82f6", "Sur": "#8b5cf6"},
                        text_auto=True
                    )
                    fig_gap.update_layout(height=280, margin=dict(t=10, b=10, l=10, r=10))
                    st.plotly_chart(fig_gap, use_container_width=True)
                    
            with col_g6:
                st.markdown("##### 🔔 Alertas de Fecha de Entrega Próxima (Semáforo)")
                df_active_warnings = df_p[df_p['status'] == 'En Proceso'].copy()
                if df_active_warnings.empty:
                    st.caption("No hay cotizaciones activas en proceso de entrega.")
                else:
                    warnings_list = []
                    for idx, row_warn in df_active_warnings.iterrows():
                        try:
                            tgt_dt = datetime.strptime(row_warn['target_date'], "%Y-%m-%d").date()
                        except:
                            continue
                            
                        days_left = (tgt_dt - date.today()).days
                        
                        # Map current stage to friendly name and responsible
                        stg = row_warn['current_stage']
                        if stg == 1:
                            stg_name, resp_name = "Paso 1: Levantamiento", row_warn['assigned_ventas']
                        elif stg == 2:
                            stg_name, resp_name = "Paso 2: Minuta de Trabajo", f"{row_warn['assigned_ventas']} / {row_warn['assigned_lider']}"
                        elif stg == 3:
                            stg_name, resp_name = "Paso 3: Catálogo de Conceptos", row_warn['assigned_lider']
                        elif stg == 4:
                            stg_name, resp_name = "Paso 4: Elaboración de Cotización", row_warn['assigned_costos']
                        elif stg == 5:
                            stg_name, resp_name = "Paso 5: Revisión de Dirección", "Noe Ortiz"
                        elif stg == 6:
                            stg_name, resp_name = "Paso 6: Entrega al Cliente", row_warn['assigned_ventas']
                        else:
                            stg_name, resp_name = "Paso 7: Cierre Comercial", "Noe Ortiz"
                            
                        warnings_list.append({
                            "id": row_warn['id'],
                            "name": row_warn['name'],
                            "days_left": days_left,
                            "stage_name": stg_name,
                            "responsible": resp_name
                        })
                    
                    if not warnings_list:
                        st.caption("No hay alertas disponibles.")
                    else:
                        # Sort by days_left ascending (most urgent first)
                        warnings_list = sorted(warnings_list, key=lambda x: x['days_left'])
                        
                        for w in warnings_list[:5]: # Show top 5 urgent
                            if w['days_left'] < 0:
                                st.error(f"🔴 **{w['id']} - {w['name']}** (VENCIDO HACE {abs(w['days_left'])} DÍAS)  \nAtorado en: *{w['stage_name']}* | Responsable: `{w['responsible']}`")
                            elif w['days_left'] <= 7:
                                st.warning(f"🟠 **{w['id']} - {w['name']}** (URGENTE, VENCE EN {w['days_left']} DÍAS)  \nAtorado en: *{w['stage_name']}* | Responsable: `{w['responsible']}`")
                            else:
                                st.success(f"🟢 **{w['id']} - {w['name']}** (En tiempo, quedan {w['days_left']} días)  \nUbicación: *{w['stage_name']}* | Responsable: `{w['responsible']}`")


# ==========================================
# MÓDULO 2: TABLERO DE PROYECTOS
# ==========================================
if "📋 Tablero de Proyectos" in tab_dict:
    with tab_dict["📋 Tablero de Proyectos"]:
        st.subheader("📋 Pipeline General de Proyectos")
        
        # Opciones de creación y gestión exclusivas para el Admin/Director y Directores
        role_clean = str(role).strip().lower()
        if role_clean in ["admin/director", "director comercial", "director de proyectos"] or "director" in role_clean:
            col_admin_actions = st.columns(2)
            with col_admin_actions[0]:
                with st.expander("➕ Crear Nuevo Registro de Licitación"):
                    with st.form("Add Project Form"):
                        p_name = st.text_input("Nombre de la Obra")
                        p_client = st.text_input("Cliente")
                        p_state = st.selectbox("Estado de la República", list(ESTADOS_MEXICO.keys()))
                        
                        # Cargar listas dinámicas de usuarios
                        conn = get_db_connection()
                        cost_users = [u['full_name'] for u in conn.execute("SELECT full_name FROM users WHERE role LIKE '%Costos%'").fetchall()]
                        sales_users = [u['full_name'] for u in conn.execute("SELECT full_name FROM users WHERE role = 'Ventas'").fetchall()]
                        conn.close()
                        
                        p_costos = st.selectbox("Analista de Costos Asignado", cost_users if cost_users else ["Lic. Roberto (Director de Costos)"])
                        
                        p_comm_responsibility = st.radio("Responsable de Levantamiento y Entrega (Pasos 1 y 6)", ["Agente de Ventas", "Líder Regional"], key="p_comm_responsibility_new")
                        if p_comm_responsibility == "Líder Regional":
                            st.info("💡 El Líder Regional de la zona asumirá el Levantamiento y la Entrega del proyecto.")
                            p_ventas = None
                        else:
                            p_ventas = st.selectbox("Agente de Ventas Responsable", sales_users if sales_users else ["Ing. Carlos"])
                            
                        p_priority = st.selectbox("Prioridad de Licitación", ["Alta", "Media", "Baja"], index=1)
                        p_target = st.date_input("Fecha Compromiso de Entrega")
                        
                        # Mostrar el código inteligente que se generará
                        temp_code = generate_next_project_id(p_state)
                        st.info(f"**Código de cotización asignado:** {temp_code}")
                        
                        btn_add_proj = st.form_submit_button("Guardar en Pipeline 📂")
                        
                        if btn_add_proj:
                            if not p_name or not p_client:
                                st.error("Todos los campos obligatorios deben ser llenados.")
                            else:
                                final_code = generate_next_project_id(p_state)
                                region_auto = ESTADOS_MEXICO[p_state]
                                zone_auto = "Sur" if "Sur" in region_auto else "Norte"
                                
                                # Mapear líder automáticamente consultando la base de datos de usuarios para esa región
                                conn_lider = get_db_connection()
                                lider_db = conn_lider.execute("SELECT full_name FROM users WHERE role = ?", (region_auto,)).fetchone()
                                conn_lider.close()
                                
                                if lider_db:
                                    assigned_leader = lider_db['full_name']
                                else:
                                    assigned_leader = region_auto
                                    
                                final_ventas = assigned_leader if p_comm_responsibility == "Líder Regional" else p_ventas
                                
                                conn = get_db_connection()
                                try:
                                    conn.execute('''
                                        INSERT INTO projects (
                                            id, name, client, total_amount, final_amount, state, zone, 
                                            assigned_lider, assigned_costos, assigned_ventas, priority, status, current_stage, 
                                            created_at, target_date
                                        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                                    ''', (
                                        final_code, p_name, p_client, 0.0, 0.0, p_state, zone_auto,
                                        assigned_leader, p_costos, final_ventas, p_priority, "En Proceso", 1,
                                        date.today().strftime("%Y-%m-%d"), p_target.strftime("%Y-%m-%d")
                                    ))
                                    conn.commit()
                                    log_audit(final_code, st.session_state.full_name, role, f"Creó licitación con código {final_code}")
                                    st.success(f"Proyecto {final_code} guardado con éxito.")
                                    st.rerun()
                                except Exception as e:
                                    st.error(f"Error al guardar: {e}")
                                finally:
                                    conn.close()
                                    
            with col_admin_actions[1]:
                with st.expander("⚙️ Gestionar / Editar / Eliminar Licitación"):
                    conn = get_db_connection()
                    all_p_db = conn.execute("SELECT * FROM projects").fetchall()
                    conn.close()
                    
                    if all_p_db:
                        p_to_del_label = st.selectbox("Seleccionar Proyecto para Operaciones", [f"{p['id']} - {p['name']}" for p in all_p_db], key="manage_select_proj")
                        sel_del_id = p_to_del_label.split(" - ")[0]
                        p_details = next((dict(row) for row in all_p_db if row['id'] == sel_del_id), None)
                        
                        if p_details:
                            st.markdown("##### ✏️ Modificar Parámetros de Licitación")
                            
                            edit_name = st.text_input("Nombre de la Obra / Proyecto", value=p_details['name'], key="edit_p_name")
                            edit_client = st.text_input("Cliente", value=p_details['client'], key="edit_p_client")
                            
                            # Ensure columns are safely accessed
                            curr_target_val = date.today()
                            if p_details.get('target_date'):
                                try:
                                    curr_target_val = datetime.strptime(p_details['target_date'], "%Y-%m-%d").date()
                                except:
                                    pass
                            edit_target = st.date_input("Fecha Compromiso de Entrega", value=curr_target_val, key="edit_p_target")
                            
                            curr_priority = p_details.get('priority', 'Media')
                            if not curr_priority:
                                curr_priority = 'Media'
                            edit_priority = st.selectbox("Prioridad", ["Alta", "Media", "Baja"], index=["Alta", "Media", "Baja"].index(curr_priority) if curr_priority in ["Alta", "Media", "Baja"] else 1, key="edit_p_priority")
                            
                            conn = get_db_connection()
                            cost_users = [u['full_name'] for u in conn.execute("SELECT full_name FROM users WHERE role LIKE '%Costos%'").fetchall()]
                            sales_users = [u['full_name'] for u in conn.execute("SELECT full_name FROM users WHERE role = 'Ventas'").fetchall()]
                            conn.close()
                            
                            curr_costos = p_details.get('assigned_costos', '')
                            edit_costos = st.selectbox("Analista de Costos Asignado", cost_users if cost_users else ["Lic. Roberto (Director de Costos)"], index=cost_users.index(curr_costos) if curr_costos in cost_users else 0, key="edit_p_costos")
                            
                            curr_ventas = p_details.get('assigned_ventas', '')
                            curr_leader = p_details.get('assigned_lider', '')
                            
                            is_lider_comm_init = 1 if curr_ventas == curr_leader else 0
                            edit_comm_responsibility = st.radio("Responsable de Levantamiento y Entrega (Pasos 1 y 6)", ["Agente de Ventas", "Líder Regional"], index=is_lider_comm_init, key="edit_p_comm_responsibility")
                            
                            if edit_comm_responsibility == "Líder Regional":
                                st.info(f"💡 El Líder Regional ({curr_leader}) asumirá la responsabilidad comercial de este proyecto.")
                                edit_ventas_val = curr_leader
                            else:
                                edit_ventas = st.selectbox("Agente de Ventas Responsable", sales_users if sales_users else ["Ing. Carlos"], index=sales_users.index(curr_ventas) if curr_ventas in sales_users else 0, key="edit_p_ventas")
                                edit_ventas_val = edit_ventas

                            st.markdown("---")
                            st.markdown("##### 🛠️ Configuración de Compuertas / Pasos Omitidos")
                            st.caption("Marque o desmarque los pasos completados manualmente. Esto le permite configurar el proyecto si ya inició o se adelantaron etapas.")
                            
                            col_gates1, col_gates2 = st.columns(2)
                            with col_gates1:
                                toggle_s1 = st.checkbox("Paso 1 Completado (Levantamiento)", value=(p_details.get('step1_completed', 0) == 1), key="chk_gate_s1")
                                toggle_s2_v = st.checkbox("Paso 2: Reunión Confirmada por Ventas", value=(p_details.get('step2_ventas_done', 0) == 1), key="chk_gate_s2_v")
                                toggle_s2_l = st.checkbox("Paso 2: Reunión Confirmada por Líder", value=(p_details.get('step2_lider_done', 0) == 1), key="chk_gate_s2_l")
                                toggle_s2 = st.checkbox("Paso 2 Completado (Minuta Cargada)", value=(p_details.get('step2_completed', 0) == 1), key="chk_gate_s2")
                            with col_gates2:
                                toggle_s3 = st.checkbox("Paso 3 Completado (Catálogo)", value=(p_details.get('step3_completed', 0) == 1), key="chk_gate_s3")
                                toggle_s4 = st.checkbox("Paso 4 Completado (Cotización)", value=(p_details.get('step4_completed', 0) == 1), key="chk_gate_s4")
                                toggle_s5 = st.checkbox("Paso 5 Completado (Revisión Dirección)", value=(p_details.get('step5_completed', 0) == 1), key="chk_gate_s5")
                                toggle_s6 = st.checkbox("Paso 6 Completado (Entrega al Cliente)", value=(p_details.get('step6_completed', 0) == 1), key="chk_gate_s6")
                                
                            edit_stage = st.selectbox("Ubicación Actual (Paso de la Compuerta Activa)", [
                                "Paso 1: Levantamiento (Ventas)",
                                "Paso 2: Minuta (Ventas & Líder)",
                                "Paso 3: Catálogo Conceptos (Líder)",
                                "Paso 4: Cotización (Costos)",
                                "Paso 5: Revisión y Aprobación (Dirección)",
                                "Paso 6: Entrega al Cliente (Ventas)",
                                "Paso 7: Cierre Comercial (Dirección)"
                            ], index=max(0, min(6, p_details.get('current_stage', 1) - 1)), key="edit_p_stage")
                            
                            edit_stage_val = [
                                "Paso 1: Levantamiento (Ventas)",
                                "Paso 2: Minuta (Ventas & Líder)",
                                "Paso 3: Catálogo Conceptos (Líder)",
                                "Paso 4: Cotización (Costos)",
                                "Paso 5: Revisión y Aprobación (Dirección)",
                                "Paso 6: Entrega al Cliente (Ventas)",
                                "Paso 7: Cierre Comercial (Dirección)"
                            ].index(edit_stage) + 1

                            if st.button("Guardar Cambios 💾", type="primary", use_container_width=True, key="save_edit_p_btn"):
                                if not edit_name or not edit_client:
                                    st.error("El nombre y el cliente son campos obligatorios.")
                                else:
                                    prev_stage = p_details.get('current_stage', 1) if p_details else 1
                                    conn = get_db_connection()
                                    conn.execute('''
                                        UPDATE projects 
                                        SET name = ?, client = ?, target_date = ?, priority = ?, assigned_costos = ?, assigned_ventas = ?,
                                            step1_completed = ?, step2_ventas_done = ?, step2_lider_done = ?, step2_completed = ?,
                                            step3_completed = ?, step4_completed = ?, step5_completed = ?, step6_completed = ?,
                                            current_stage = ?
                                        WHERE id = ?
                                    ''', (
                                        edit_name, edit_client, edit_target.strftime("%Y-%m-%d"), edit_priority, edit_costos, edit_ventas_val,
                                        1 if toggle_s1 else 0, 1 if toggle_s2_v else 0, 1 if toggle_s2_l else 0, 1 if toggle_s2 else 0,
                                        1 if toggle_s3 else 0, 1 if toggle_s4 else 0, 1 if toggle_s5 else 0, 1 if toggle_s6 else 0,
                                        edit_stage_val, sel_del_id
                                    ))
                                    conn.commit()
                                    conn.close()
                                    
                                    # Trigger notifications if stage advanced
                                    if edit_stage_val > prev_stage:
                                        dispatch_step_completion_notifications(sel_del_id, edit_stage_val - 1)
                                    
                                    # Invalidate cache
                                    try:
                                        get_cached_dashboard_stats.clear()
                                    except:
                                        pass
                                        
                                    log_audit(sel_del_id, st.session_state.full_name, role, f"Modificó licitación {sel_del_id} (Prioridad: {edit_priority}, Fecha de entrega: {edit_target.strftime('%Y-%m-%d')})")
                                    st.success(f"Cambios en proyecto {sel_del_id} guardados con éxito.")
                                    st.rerun()
                                    
                            st.markdown("---")
                            st.markdown("##### 🚨 Zona de Peligro")
                            confirm_del = st.checkbox("Confirmo que deseo ELIMINAR permanentemente esta cotización y todos sus documentos.", key="confirm_del_check")
                            if st.button("Eliminar permanentemente 🗑️", type="primary", disabled=not confirm_del, key="delete_p_btn"):
                                conn = get_db_connection()
                                conn.execute("DELETE FROM projects WHERE id = ?", (sel_del_id,))
                                
                                # Borrar archivos físicos
                                uploaded_files = conn.execute("SELECT file_path FROM uploads WHERE project_id = ?", (sel_del_id,)).fetchall()
                                for f in uploaded_files:
                                    if f['file_path'] and os.path.exists(f['file_path']):
                                        try:
                                            os.remove(f['file_path'])
                                        except:
                                            pass
                                conn.execute("DELETE FROM uploads WHERE project_id = ?", (sel_del_id,))
                                conn.commit()
                                conn.close()
                                
                                # Invalidate cache
                                try:
                                    get_cached_dashboard_stats.clear()
                                except:
                                    pass
                                    
                                log_audit(sel_del_id, st.session_state.full_name, role, "Eliminó licitación y todos sus archivos asociados.")
                                st.success(f"Proyecto {sel_del_id} eliminado.")
                                st.rerun()
                    else:
                        st.caption("No hay proyectos en base de datos")

        # Cargar y desplegar lista nativa de proyectos en Pipeline
        conn = get_db_connection()
        proyectos = conn.execute("SELECT * FROM projects").fetchall()
        conn.close()
        proyectos = filter_user_projects([dict(r) for r in proyectos])
        
        if not proyectos:
            st.warning("No hay proyectos registrados en este momento.")
        else:
            df_projs = pd.DataFrame([dict(p) for p in proyectos])
            # Ensure all expected columns are present in df_projs
            expected_db_cols = ['id', 'name', 'client', 'total_amount', 'final_amount', 'state', 'zone', 'assigned_lider', 'assigned_costos', 'assigned_ventas', 'assigned_ventas_2', 'priority', 'status', 'current_stage']
            for col in expected_db_cols:
                if col not in df_projs.columns:
                    df_projs[col] = 0.0 if col in ['total_amount', 'final_amount'] else None
            
            # Sort by Priority categories
            df_projs['priority_sort'] = pd.Categorical(df_projs['priority'].fillna('Media'), categories=['Alta', 'Media', 'Baja'], ordered=True)
            df_projs = df_projs.sort_values(by=['priority_sort', 'id'], ascending=[True, False]).drop(columns=['priority_sort'])
            
            df_display = df_projs.rename(columns={
                'id': 'ID Proyecto',
                'name': 'Obra / Proyecto',
                'client': 'Cliente',
                'final_amount': 'Monto Cotizado ($)',
                'state': 'Estado',
                'zone': 'Zona',
                'assigned_lider': 'Líder Regional',
                'assigned_costos': 'Analista de Costos',
                'assigned_ventas': 'Agente de Ventas',
                'assigned_ventas_2': 'Agente Ventas 2',
                'priority': 'Prioridad',
                'status': 'Estatus Comercial',
                'current_stage': 'Paso Actual'
            })
            
            # Map paso actual to descriptive text
            steps_desc = {
                1: "1. Levantamiento (Ventas)",
                2: "2. Minuta (Ventas & Líder)",
                3: "3. Catálogo Conceptos (Líder)",
                4: "4. Cotización (Costos)",
                5: "5. Revisión y Aprobación (Dirección)",
                6: "6. Entrega al Cliente (Ventas)",
                7: "7. Cierre Comercial (Dirección)"
            }
            df_display['Paso Actual'] = df_display['Paso Actual'].map(steps_desc)
            
            cols_order_display = ['ID Proyecto', 'Obra / Proyecto', 'Cliente', 'Monto Cotizado ($)', 'Estado', 'Zona', 'Agente de Ventas', 'Líder Regional', 'Analista de Costos', 'Prioridad', 'Paso Actual', 'Estatus Comercial']
            st.dataframe(
                df_display[cols_order_display],
                use_container_width=True,
                hide_index=True,
                column_config={
                    "Monto Cotizado ($)": st.column_config.NumberColumn(format="$%,.2f"),
                }
            )

# ==========================================
# MÓDULO 3: COMPUERTAS TÉCNICAS (FLUJO SECUENCIAL)
# ==========================================
if "✔️ Compuertas Técnicas" in tab_dict:
    with tab_dict["✔️ Compuertas Técnicas"]:
        st.subheader("✔️ Seguimiento de Pasos de Cotización")
        st.write("Estructura secuencial obligatoria para cada proyecto. Complete cada tarea para desbloquear el paso siguiente:")
        
        conn = get_db_connection()
        projs = conn.execute("SELECT * FROM projects").fetchall()
        conn.close()
        projs = filter_user_projects([dict(r) for r in projs])
        
        if not projs:
            st.info("No hay proyectos registrados para procesar.")
        else:
            # Selector de proyectos
            proj_dict = {f"{p['id']} - {p['name']}": p for p in projs}
            sel_proj_label = st.selectbox("📂 Seleccione la Obra / Proyecto", list(proj_dict.keys()), key="seq_proj_select")
            p = proj_dict[sel_proj_label]
            
            # Tarjetas informativas superiores
            col_met1, col_met2, col_met3, col_met4, col_met5 = st.columns(5)
            with col_met1:
                st.metric("Agente de Ventas", p['assigned_ventas'])
            with col_met2:
                st.metric("Líder Regional", p['assigned_lider'])
            with col_met3:
                st.metric("Analista de Costos", p['assigned_costos'])
            with col_met4:
                st.metric("Monto Cotizado ($)", f"${p['final_amount']:,.2f}")
            with col_met5:
                st.metric("Estatus Licitación", p['status'])
                
            st.markdown(f"**Prioridad:** `{p.get('priority', 'Media') if p.get('priority') else 'Media'}`  |  **Fecha Límite de Entrega Comercial:** `{p.get('target_date', 'No definida')}`")
                
            st.markdown("<br>", unsafe_allow_html=True)
            report_bytes = generate_docx_report(p['id'])
            if report_bytes:
                col_d_dl, col_d_save = st.columns(2)
                with col_d_dl:
                    st.download_button(
                        label=f"📥 Descargar Dossier (Web)",
                        data=report_bytes,
                        file_name=f"Dossier_Ejecutivo_{p['id']}.docx",
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                        use_container_width=True
                    )
                with col_d_save:
                    if st.button(f"💾 Guardar en Escritorio (Word)", key=f"save_pc_dossier_{p['id']}", use_container_width=True):
                        save_file_directly_to_pc(report_bytes, f"Dossier_Ejecutivo_{p['id']}.docx")
                
            st.markdown("---")
            
            # Renderizado de pestañas de flujo paso a paso
            step_tabs = st.tabs([
                "Paso 1: Levantamiento 📋",
                "Paso 2: Minuta Trabajo 🤝",
                "Paso 3: Catálogo Conceptos ⚙️",
                "Paso 4: Cotización 📊",
                "Paso 5: Revisión Dirección 🔑",
                "Paso 6: Entrega Cliente 🚚",
                "Paso 7: Cierre Comercial 🏁"
            ])
            
            # Helper: Consultar archivos adjuntos para un paso específico
            def get_step_files(proj_id, step_name):
                conn = get_db_connection()
                files = conn.execute("SELECT * FROM uploads WHERE project_id = ? AND step_name = ?", (proj_id, step_name)).fetchall()
                conn.close()
                return files
            
            # Helper: Desplegar interfaz de archivos con permisos de eliminación nativos
            def display_files_interface(proj_id, step_name, active_user, is_readonly=False):
                st.markdown("**📂 Documentación y Archivos Adjuntos**")
                files = get_step_files(proj_id, step_name)
                if not files:
                    st.caption("No se han cargado documentos en este paso.")
                else:
                    for f in files:
                        col_file_name, col_file_dl, col_file_del = st.columns([3, 2, 1])
                        with col_file_name:
                            st.write(f"📄 {f['filename']} (Subido por: {f['uploaded_by']})")
                        with col_file_dl:
                            # Descargar archivo (Global DB-driven fallback to local disk)
                            f_content = None
                            if 'file_data' in dict(f) and f['file_data'] is not None:
                                try:
                                    f_content = bytes(f['file_data'])
                                except Exception:
                                    pass
                            if f_content is None:
                                try:
                                    with open(f['file_path'], "rb") as file_bytes:
                                        f_content = file_bytes.read()
                                except:
                                    pass
                                    
                            if f_content is not None:
                                col_inner_dl, col_inner_save = st.columns(2)
                                with col_inner_dl:
                                    st.download_button(
                                        label="Web 📥",
                                        data=f_content,
                                        file_name=f['filename'],
                                        key=f"dl_{f['id']}",
                                        use_container_width=True
                                    )
                                with col_inner_save:
                                    if st.button("PC 💾", key=f"save_pc_att_{f['id']}", use_container_width=True):
                                        save_file_directly_to_pc(f_content, f['filename'])
                            else:
                                st.error("Archivo no disponible")
                        with col_file_del:
                            # Eliminar archivo (Solo si es quien lo cargó o Admin, y no es visualizador)
                            if not is_readonly and (active_user == f['uploaded_by'] or role == "Admin/Director"):
                                if st.button("Quitar 🗑️", key=f"del_{f['id']}"):
                                    conn = get_db_connection()
                                    conn.execute("DELETE FROM uploads WHERE id = ?", (f['id'],))
                                    conn.commit()
                                    conn.close()
                                    if os.path.exists(f['file_path']):
                                        try:
                                            os.remove(f['file_path'])
                                        except:
                                            pass
                                    log_audit(proj_id, st.session_state.full_name, role, f"Eliminó archivo {f['filename']} del paso {step_name}")
                                    st.success("Archivo eliminado.")
                                    st.rerun()
                                    
            # Permiso lectura general
            is_readonly = (role == "Ingeniero")
            
            # ---------------------------------------------
            # PASO 1: LEVANTAMIENTO (Ventas)
            # ---------------------------------------------
            with step_tabs[0]:
                st.markdown("### Paso 1: Crear Levantamiento Técnico")
                if p['assigned_ventas'] == p['assigned_lider']:
                    st.write(f"**Asignado a:** Líder Regional - *{p['assigned_ventas']}* (Asignación Directa)")
                else:
                    st.write(f"**Asignado a:** Agente de Ventas - *{p['assigned_ventas']}*")
                
                # Desplegar Archivos
                display_files_interface(p['id'], "step1_levantamiento", st.session_state.user_name, is_readonly)
                
                # Cargar Archivo (Solo Ventas asignado o Admin)
                is_authorized_s1 = (st.session_state.full_name == p['assigned_ventas'] or role == "Admin/Director") and not is_readonly
                
                if p['step1_completed'] == 0:
                    st.info("Estatus: Esperando carga de levantamiento técnico por el agente de ventas.")
                    if is_authorized_s1:
                        uploaded_file_s1 = st.file_uploader("Cargar evidencia de levantamiento (PDF, Imagen, Word, etc.)", key="uploader_s1")
                        if uploaded_file_s1:
                            if st.button("Guardar Evidencia y Validar Paso ✔️", key="btn_s1"):
                                file_path = save_uploaded_file(uploaded_file_s1, p['id'], "step1_levantamiento")
                                conn = get_db_connection()
                                conn.execute("UPDATE projects SET step1_completed = 1, current_stage = 2 WHERE id = ?", (p['id'],))
                                conn.commit()
                                conn.close()
                                dispatch_step_completion_notifications(p['id'], 1)
                                log_audit(p['id'], st.session_state.full_name, role, "Completó Paso 1: Carga de levantamiento")
                                st.success("Paso 1 completado con éxito. Paso 2 desbloqueado.")
                                st.rerun()
                else:
                    st.success("✔️ Paso 1 Completado: Evidencia de levantamiento cargada.")
                    if is_authorized_s1:
                        # Opción de re-subir archivos adicionales
                        uploaded_file_s1_extra = st.file_uploader("Cargar archivo adicional", key="uploader_s1_extra")
                        if uploaded_file_s1_extra:
                            if st.button("Subir archivo adicional", key="btn_s1_extra"):
                                file_path = save_uploaded_file(uploaded_file_s1_extra, p['id'], "step1_levantamiento")
                                st.success("Archivo subido con éxito.")
                                st.rerun()

            # ---------------------------------------------
            # PASO 2: MINUTA TRABAJO (Ventas & Líder)
            # ---------------------------------------------
            with step_tabs[1]:
                st.markdown("### Paso 2: Reunión de Seguimiento y Minuta de Trabajo")
                
                is_lider_commercial = (p['assigned_ventas'] == p['assigned_lider'])
                if is_lider_commercial:
                    st.write(f"**Ventas & Líder Regional Responsable (Mismo Rol):** {p['assigned_lider']}")
                else:
                    st.write(f"**Ventas Responsable:** {p['assigned_ventas']}")
                    st.write(f"**Líder Regional Responsable:** {p['assigned_lider']}")
                
                display_files_interface(p['id'], "step2_minuta", st.session_state.user_name, is_readonly)
                
                # Validar si el paso previo está completado
                if p['step1_completed'] == 0:
                    st.warning("🔒 Este paso se encuentra bloqueado. Complete el Paso 1 para poder acceder.")
                else:
                    is_ventas = (st.session_state.full_name == p['assigned_ventas'] or role == "Admin/Director") and not is_readonly
                    is_lider = (st.session_state.full_name == p['assigned_lider'] or st.session_state.user_role == p['assigned_lider'] or role == "Admin/Director") and not is_readonly
                    
                    if is_lider_commercial:
                        st.markdown("##### ☑️ Confirmación de Reunión Realizada")
                        if p['step2_lider_done'] == 1:
                            st.success("✔️ Líder Regional (Responsable Único): Reunión Confirmada")
                        else:
                            st.warning("⏳ Líder Regional: Reunión Pendiente")
                            if is_lider:
                                if st.button("Confirmar Reunión 🤝", key="btn_confirm_l_s2_single", use_container_width=True):
                                    conn = get_db_connection()
                                    conn.execute("UPDATE projects SET step2_lider_done = 1, step2_ventas_done = 1 WHERE id = ?", (p['id'],))
                                    conn.commit()
                                    conn.close()
                                    st.success("Reunión confirmada.")
                                    st.rerun()
                    else:
                        st.markdown("##### ☑️ Confirmación de Reunión Realizada (Doble Check)")
                        col_chk1, col_col2 = st.columns(2)
                        with col_chk1:
                            if p['step2_ventas_done'] == 1:
                                st.success("✔️ Ventas: Reunión Confirmada")
                            else:
                                st.warning("⏳ Ventas: Reunión Pendiente")
                                if is_ventas:
                                    if st.button("Confirmar Reunión (Ventas) 🤝", key="btn_confirm_v_s2", use_container_width=True):
                                        conn = get_db_connection()
                                        conn.execute("UPDATE projects SET step2_ventas_done = 1 WHERE id = ?", (p['id'],))
                                        conn.commit()
                                        conn.close()
                                        st.success("Reunión confirmada por Ventas.")
                                        st.rerun()
                        with col_col2:
                            if p['step2_lider_done'] == 1:
                                st.success("✔️ Líder Regional: Reunión Confirmada")
                            else:
                                st.warning("⏳ Líder Regional: Reunión Pendiente")
                                if is_lider:
                                    if st.button("Confirmar Reunión (Líder) 🤝", key="btn_confirm_l_s2", use_container_width=True):
                                        conn = get_db_connection()
                                        conn.execute("UPDATE projects SET step2_lider_done = 1 WHERE id = ?", (p['id'],))
                                        conn.commit()
                                        conn.close()
                                        st.success("Reunión confirmada por el Líder Regional.")
                                        st.rerun()
                            
                    # Carga de minuta
                    if p['step2_completed'] == 0:
                        if p['step2_ventas_done'] == 1 and p['step2_lider_done'] == 1:
                            st.info("Ambas partes han confirmado la reunión. Proceda a cargar la minuta de trabajo firmada para validar la compuerta:")
                            if is_lider:
                                uploaded_file_s2 = st.file_uploader("Cargar archivo de minuta de trabajo", key="uploader_s2")
                                if uploaded_file_s2:
                                    if st.button("Guardar Minuta y Validar Paso ✔️", key="btn_s2", use_container_width=True):
                                        file_path = save_uploaded_file(uploaded_file_s2, p['id'], "step2_minuta")
                                        conn = get_db_connection()
                                        conn.execute("UPDATE projects SET step2_completed = 1, current_stage = 3 WHERE id = ?", (p['id'],))
                                        conn.commit()
                                        conn.close()
                                        dispatch_step_completion_notifications(p['id'], 2)
                                        log_audit(p['id'], st.session_state.full_name, role, "Completó Paso 2: Carga de minuta de trabajo")
                                        st.success("Paso 2 completado. Paso 3 desbloqueado.")
                                        st.rerun()
                            elif is_ventas and not is_lider_commercial:
                                st.warning("📢 Esperando que el Líder Regional asignado cargue el archivo de la minuta de trabajo firmada.")
                        else:
                            st.warning("Esperando confirmación de 'Reunión hecha' para habilitar la carga de documentos.")
                    else:
                        st.success("✔️ Paso 2 Completado: Minuta cargada y validada.")
                        if is_lider:
                            uploaded_file_s2_extra = st.file_uploader("Cargar archivo de minuta adicional", key="uploader_s2_extra")
                            if uploaded_file_s2_extra:
                                if st.button("Subir minuta adicional", key="btn_s2_extra"):
                                    file_path = save_uploaded_file(uploaded_file_s2_extra, p['id'], "step2_minuta")
                                    st.success("Archivo subido con éxito.")
                                    st.rerun()
                        elif is_ventas and not is_lider_commercial:
                            st.info("📢 Solo el Líder Regional asignado o la Dirección pueden subir minutas de trabajo adicionales.")

            # ---------------------------------------------
            # PASO 3: CATÁLOGO CONCEPTOS (Líder)
            # ---------------------------------------------
            with step_tabs[2]:
                st.markdown("### Paso 3: Catálogo de Conceptos Técnico")
                st.write(f"**Asignado a:** Líder Regional - *{p['assigned_lider']}*")
                
                display_files_interface(p['id'], "step3_catalogo", st.session_state.user_name, is_readonly)
                
                if p['step2_completed'] == 0:
                    st.warning("🔒 Este paso se encuentra bloqueado. Complete el Paso 2 para poder acceder.")
                else:
                    is_authorized_s3 = (st.session_state.full_name == p['assigned_lider'] or st.session_state.user_role == p['assigned_lider'] or role == "Admin/Director") and not is_readonly
                    
                    if p['step3_completed'] == 0:
                        st.info("Estatus: Esperando catálogo de conceptos técnico por el Líder Regional.")
                        if is_authorized_s3:
                            uploaded_file_s3 = st.file_uploader("Cargar Catálogo de Conceptos (Excel, PDF)", key="uploader_s3")
                            if uploaded_file_s3:
                                if st.button("Guardar Catálogo y Validar Paso ✔️", key="btn_s3"):
                                    file_path = save_uploaded_file(uploaded_file_s3, p['id'], "step3_catalogo")
                                    conn = get_db_connection()
                                    conn.execute("UPDATE projects SET step3_completed = 1, current_stage = 4 WHERE id = ?", (p['id'],))
                                    conn.commit()
                                    conn.close()
                                    dispatch_step_completion_notifications(p['id'], 3)
                                    log_audit(p['id'], st.session_state.full_name, role, "Completó Paso 3: Carga de catálogo de conceptos")
                                    st.success("Paso 3 completado. Paso 4 desbloqueado.")
                                    st.rerun()
                    else:
                        st.success("✔️ Paso 3 Completado: Catálogo de conceptos técnico cargado.")
                        if is_authorized_s3:
                            uploaded_file_s3_extra = st.file_uploader("Cargar archivo técnico adicional", key="uploader_s3_extra")
                            if uploaded_file_s3_extra:
                                if st.button("Subir archivo adicional", key="btn_s3_extra"):
                                    file_path = save_uploaded_file(uploaded_file_s3_extra, p['id'], "step3_catalogo")
                                    st.success("Archivo subido con éxito.")
                                    st.rerun()

            # ---------------------------------------------
            # PASO 4: COTIZACIÓN (Costos)
            # ---------------------------------------------
            with step_tabs[3]:
                st.markdown("### Paso 4: Elaboración de Cotización de Precios")
                st.write(f"**Asignado a:** Analista de Costos - *{p['assigned_costos']}*")
                
                display_files_interface(p['id'], "step4_cotizacion", st.session_state.user_name, is_readonly)
                
                if p['step3_completed'] == 0:
                    st.warning("🔒 Este paso se encuentra bloqueado. Complete el Paso 3 para poder acceder.")
                else:
                    is_authorized_s4 = (st.session_state.full_name == p['assigned_costos'] or st.session_state.user_role == p['assigned_costos'] or role == "Admin/Director") and not is_readonly
                    
                    # El analista de costos ve el catálogo y puede validarlo con checks interactivos nativos
                    st.markdown("##### 📝 Verificación de Conceptos Técnicos")
                    chk_c1 = st.checkbox("Conceptos alineados con catálogo", key="chk_c1", disabled=not is_authorized_s4)
                    chk_c2 = st.checkbox("Análisis de precios unitarios formulado", key="chk_c2", disabled=not is_authorized_s4)
                    chk_c3 = st.checkbox("Márgenes comerciales integrados", key="chk_c3", disabled=not is_authorized_s4)
                    
                    if p['step4_completed'] == 0:
                        if is_authorized_s4:
                            st.markdown("##### 📤 Cargar Cotización de Precios")
                            uploaded_file_s4 = st.file_uploader("Cargar cotización final de precios", key="uploader_s4")
                            
                            # Validar que los checks se marquen antes de completar
                            if uploaded_file_s4:
                                if chk_c1 and chk_c2 and chk_c3:
                                    if st.button("Cargar Cotización y Validar Paso ✔️", key="btn_s4"):
                                        file_path = save_uploaded_file(uploaded_file_s4, p['id'], "step4_cotizacion")
                                        conn = get_db_connection()
                                        conn.execute("UPDATE projects SET step4_completed = 1, current_stage = 5 WHERE id = ?", (p['id'],))
                                        conn.commit()
                                        conn.close()
                                        dispatch_step_completion_notifications(p['id'], 4)
                                        log_audit(p['id'], st.session_state.full_name, role, "Completó Paso 4: Elaboró y cargó cotización")
                                        st.success("Paso 4 completado con éxito. Paso 5 desbloqueado para revisión de dirección.")
                                        st.rerun()
                                else:
                                    st.warning("Debe marcar las 3 casillas de verificación técnica para poder validar y guardar la cotización.")
                    else:
                        st.success("✔️ Paso 4 Completado: Cotización final cargada y verificada.")
                        if is_authorized_s4:
                            uploaded_file_s4_extra = st.file_uploader("Cargar archivo de cotización adicional", key="uploader_s4_extra")
                            if uploaded_file_s4_extra:
                                if st.button("Subir cotización adicional", key="btn_s4_extra"):
                                    file_path = save_uploaded_file(uploaded_file_s4_extra, p['id'], "step4_cotizacion")
                                    st.success("Archivo subido con éxito.")
                                    st.rerun()

            # ---------------------------------------------
            # PASO 5: REVISIÓN DIRECCIÓN (Admin / Director)
            # ---------------------------------------------
            with step_tabs[4]:
                st.markdown("### Paso 5: Revisión de Cotización y Aprobación de Costos")
                st.write("**Asignado a:** Dirección General / Director Comercial / Director de Proyectos")
                
                if p['step4_completed'] == 0:
                    st.warning("🔒 Este paso se encuentra bloqueado. Complete el Paso 4 para poder acceder.")
                else:
                    role_clean = str(role).strip().lower()
                    is_authorized_s5 = (role_clean in ["admin/director", "director comercial", "director de proyectos"] or "director" in role_clean) and not is_readonly
                    
                    # Director puede ver y descargar los archivos cargados por el analista
                    st.markdown("##### 📥 Descargar Licitación Propuesta")
                    files_proposal = get_step_files(p['id'], "step4_cotizacion")
                    if not files_proposal:
                        st.caption("No hay propuesta cargada.")
                    else:
                        for f in files_proposal:
                            # Global DB-driven fallback to local disk
                            f_content = None
                            if 'file_data' in dict(f) and f['file_data'] is not None:
                                try:
                                    f_content = bytes(f['file_data'])
                                except Exception:
                                    pass
                            if f_content is None:
                                try:
                                    with open(f['file_path'], "rb") as file_bytes:
                                        f_content = file_bytes.read()
                                except:
                                    pass
                                    
                            if f_content is not None:
                                col_prop_dl, col_prop_save = st.columns(2)
                                with col_prop_dl:
                                    st.download_button(
                                        label=f"Descargar (Web): {f['filename']} 📥",
                                        data=f_content,
                                        file_name=f['filename'],
                                        key=f"proposal_dl_{f['id']}",
                                        use_container_width=True
                                    )
                                with col_prop_save:
                                    if st.button(f"Guardar en PC 💾: {f['filename']}", key=f"proposal_save_{f['id']}", use_container_width=True):
                                        save_file_directly_to_pc(f_content, f['filename'])
                            else:
                                st.caption(f"⚠️ {f['filename']} (No disponible)")
                                
                    st.markdown("---")
                    
                    if p['step5_completed'] == 0:
                        if is_authorized_s5:
                            st.write("Revise a detalle la cotización propuesta en costo, alcance y margen. Si es correcta, haga clic en el botón de abajo para aprobar y firmar:")
                            if st.button("Aprobar y Autorizar Cotización ✔️", key="btn_s5", type="primary", use_container_width=True):
                                conn = get_db_connection()
                                conn.execute("UPDATE projects SET step5_completed = 1, current_stage = 6 WHERE id = ?", (p['id'],))
                                conn.commit()
                                conn.close()
                                dispatch_step_completion_notifications(p['id'], 5)
                                log_audit(p['id'], st.session_state.full_name, role, "Aprobó y autorizó cotización")
                                st.success("Cotización aprobada por Dirección con éxito. Paso 6 asignado para entrega comercial.")
                                st.rerun()
                        else:
                            st.info("Estatus: Esperando revisión y firma de la Dirección o del Director Comercial/Proyectos.")
                    else:
                        st.success("✔️ Paso 5 Completado: Cotización aprobada oficialmente por Dirección.")

            # ---------------------------------------------
            # PASO 6: ENTREGA CLIENTE (Ventas)
            # ---------------------------------------------
            with step_tabs[5]:
                st.markdown("### Paso 6: Entrega de Cotización al Cliente")
                if p['assigned_ventas'] == p['assigned_lider']:
                    st.write(f"**Asignado a:** Líder Regional - *{p['assigned_ventas']}* (Asignación Directa)")
                else:
                    st.write(f"**Asignado a:** Agente de Ventas - *{p['assigned_ventas']}*")
                
                display_files_interface(p['id'], "step6_entrega", st.session_state.user_name, is_readonly)
                
                if p['step5_completed'] == 0:
                    st.warning("🔒 Este paso se encuentra bloqueado. Complete el Paso 5 para poder acceder.")
                else:
                    is_authorized_s6 = (st.session_state.full_name == p['assigned_ventas'] or st.session_state.user_role == p['assigned_ventas'] or role == "Admin/Director") and not is_readonly
                    
                    if p['step6_completed'] == 0:
                        st.info("Estatus: Esperando entrega formal al cliente y registro de comentarios finales.")
                        if is_authorized_s6:
                            with st.form("Entrega Cliente Form"):
                                final_val = st.number_input("Monto final en el que se entregó la cotización ($)", min_value=0.0, value=p['total_amount'], step=1000.0)
                                comments_s6 = st.text_area("Comentarios de entrega y observaciones del cliente")
                                chk_delivered = st.checkbox("Confirmo que la cotización ha sido entregada de forma oficial al cliente.")
                                
                                uploaded_evidence_s6 = st.file_uploader("Cargar acuse o evidencia de entrega (Opcional)", key="uploader_s6")
                                
                                btn_s6 = st.form_submit_button("Guardar Datos de Entrega ✔️")
                                if btn_s6:
                                    if not chk_delivered:
                                        st.error("Debe marcar la casilla de confirmación de entrega.")
                                    else:
                                        # Guardar evidencia si existe
                                        if uploaded_evidence_s6:
                                            save_uploaded_file(uploaded_evidence_s6, p['id'], "step6_entrega")
                                                
                                        conn = get_db_connection()
                                        conn.execute('''
                                            UPDATE projects 
                                            SET final_amount = ?, step6_completed = 1, current_stage = 7, lose_reason = ?
                                            WHERE id = ?
                                        ''', (final_val, comments_s6, p['id']))
                                        conn.commit()
                                        conn.close()
                                        dispatch_step_completion_notifications(p['id'], 6)
                                        log_audit(p['id'], st.session_state.full_name, role, f"Completó entrega de cotización por monto final ${final_val:,.2f}")
                                        st.success("Datos de entrega guardados. Paso 7 habilitado para cierre de proyecto.")
                                        st.rerun()
                    else:
                        st.success("✔️ Paso 6 Completado: Cotización entregada al cliente.")
                        st.write(f"**Monto final de entrega:** ${p['final_amount']:,.2f}")
                        st.write(f"**Observaciones:** {p['lose_reason'] or 'Sin comentarios registrados'}")

            # ---------------------------------------------
            # PASO 7: CIERRE COMERCIAL (Admin / Director)
            # ---------------------------------------------
            with step_tabs[6]:
                st.markdown("### Paso 7: Cierre Comercial de Licitación")
                st.write("**Asignado a:** Dirección General / Director Comercial / Director de Proyectos")
                
                if p['step6_completed'] == 0:
                    st.warning("🔒 Este paso se encuentra bloqueado. Complete el Paso 6 para poder acceder.")
                else:
                    role_clean = str(role).strip().lower()
                    is_authorized_s7 = (role_clean in ["admin/director", "director comercial", "director de proyectos", "director general"] or "director" in role_clean) and not is_readonly
                    
                    if p['status'] == "En Proceso":
                        if is_authorized_s7:
                            st.write("Especifique el resultado de la licitación comercial:")
                            with st.form("Cierre Licitación Form"):
                                final_res = st.selectbox("Estatus Final", ["Ganado", "Perdido", "Cancelado"])
                                gap_p = st.number_input("Porcentaje de Desfase de Precio / Margen (%)", min_value=0.0, step=1.0)
                                reason_close = st.text_area("Razón o motivo del resultado (¿Por qué se ganó o perdió?)")
                                
                                btn_s7 = st.form_submit_button("Proceder con el Cierre Oficial 🏁")
                                if btn_s7:
                                    conn = get_db_connection()
                                    conn.execute('''
                                        UPDATE projects 
                                        SET status = ?, lose_percentage_gap = ?, lose_reason = ?
                                        WHERE id = ?
                                    ''', (final_res, gap_p, reason_close, p['id']))
                                    conn.commit()
                                    conn.close()
                                    log_audit(p['id'], st.session_state.full_name, role, f"Cerró licitación comercial como '{final_res}'")
                                    st.success(f"Licitación cerrada oficialmente como '{final_res}'.")
                                    st.rerun()
                        else:
                            st.info("Estatus: Esperando cierre comercial definitivo por el Director General.")
                    else:
                        st.success(f"🏁 Licitación Cerrada: **{p['status'].upper()}**")
                        st.write(f"**Motivos:** {p['lose_reason'] or 'No especificado'}")
                        st.write(f"**Desfase de Costo:** {p['lose_percentage_gap']}%")

# ==========================================
# MÓDULO 4: KANBAN VISUAL
# ==========================================
if "🗺️ Kanban Visual" in tab_dict:
    with tab_dict["🗺️ Kanban Visual"]:
        st.subheader("🗺️ Tablero Comercial de Cotizaciones")
        
        conn = get_db_connection()
        projects_k = conn.execute("SELECT * FROM projects").fetchall()
        conn.close()
        projects_k = filter_user_projects([dict(r) for r in projects_k])
        
        col_kp, col_kg, col_kd = st.columns(3)
        
        # Helper: Renderizar tarjetas Kanban con archivos descargables integrados
        def render_kanban_card(p):
            with st.container(border=True):
                prio = p.get('priority', 'Media')
                if not prio:
                    prio = 'Media'
                color_prio_map = {"Alta": "red", "Media": "orange", "Baja": "green"}
                color_badge = color_prio_map.get(prio, "orange")
                
                st.markdown(f"**{p['id']} - {p['name']}**")
                st.markdown(f"⚠️ **Prioridad:** :{color_badge}[{prio}]")
                st.write(f"💼 **Cliente:** {p['client']}")
                st.write(f"💰 **Monto Cotizado:** ${p['total_amount']:,.2f}")
                
                ventas_team = p.get('assigned_ventas', 'Sin asignar')
                st.write(f"👤 **Ventas:** {ventas_team}")
                
                # Paso actual descriptivo
                steps_desc = {
                    1: "Paso 1: Levantamiento (Ventas)",
                    2: "Paso 2: Minuta Trabajo (Ventas & Líder)",
                    3: "Paso 3: Catálogo Conceptos (Líder)",
                    4: "Paso 4: Cotización (Costos)",
                    5: "Paso 5: Revisión Dirección (Admin/Director)",
                    6: "Paso 6: Entrega Cliente (Ventas)",
                    7: "Paso 7: Cierre Comercial (Admin/Director)"
                }
                current_step_name = steps_desc.get(p['current_stage'], "Completado")
                st.write(f"📋 **Estatus Actual:** {current_step_name}")
                
                # Barra de avance según el paso actual
                progress_val = p['current_stage'] / 7.0
                st.progress(progress_val)
                st.caption(f"Avance: Paso {p['current_stage']} de 7")
                
                # Archivos para este proyecto
                conn = get_db_connection()
                files_proj = conn.execute("SELECT * FROM uploads WHERE project_id = ?", (p['id'],)).fetchall()
                conn.close()
                
                if files_proj:
                    st.markdown("**📂 Documentos Disponibles:**")
                    for f in files_proj:
                        # Global DB-driven fallback to local disk
                        f_content = None
                        if 'file_data' in dict(f) and f['file_data'] is not None:
                            try:
                                f_content = bytes(f['file_data'])
                            except Exception:
                                pass
                        if f_content is None:
                            try:
                                with open(f['file_path'], "rb") as file_bytes:
                                    f_content = file_bytes.read()
                            except:
                                pass
                                
                        if f_content is not None:
                            col_k_dl, col_k_sv = st.columns(2)
                            with col_k_dl:
                                st.download_button(
                                    label=f"📥 {f['filename']} (Web)",
                                    data=f_content,
                                    file_name=f['filename'],
                                    key=f"kanban_dl_{p['id']}_{f['id']}",
                                    use_container_width=True
                                )
                            with col_k_sv:
                                if st.button(f"💾 Guardar {f['filename']} en PC", key=f"kanban_sv_{p['id']}_{f['id']}", use_container_width=True):
                                    save_file_directly_to_pc(f_content, f['filename'])
                        else:
                            st.caption(f"⚠️ {f['filename']} (No disponible)")
                else:
                    st.caption("Sin documentos cargados.")
        
        with col_kp:
            st.info("⏳ EN PROCESO")
            for p in projects_k:
                if p['status'] == "En Proceso":
                    render_kanban_card(p)
                        
        with col_kg:
            st.success("✔️ GANADOS")
            for p in projects_k:
                if p['status'] == "Ganado":
                    render_kanban_card(p)
                        
        with col_kd:
            st.error("🚨 PERDIDOS / CANCELADOS")
            for p in projects_k:
                if p['status'] in ["Perdido", "Cancelado"]:
                    render_kanban_card(p)

# ==========================================
# MÓDULO 5: USUARIOS Y SEGURIDAD (Exclusivo Admin)
# ==========================================
if "👥 Usuarios y Seguridad" in tab_dict:
    with tab_dict["👥 Usuarios y Seguridad"]:
        st.subheader("👥 Configuración de Usuarios e Involucrados")
        
        # --- SECCIÓN EDITAR MI PERFIL (PROGRESIVO CON CLAVE 1604) ---
        st.markdown("##### 📝 Mis Datos de Perfil")
        conn_p = get_db_connection()
        curr_user = conn_p.execute("SELECT * FROM users WHERE username = ?", (st.session_state.user_name,)).fetchone()
        conn_p.close()
        
        if curr_user:
            with st.expander("📝 Editar mi información (Nombre, Correo, Contraseña de Acceso)", expanded=False):
                with st.form("Edit My Profile Form"):
                    new_full_name = st.text_input("Mi Nombre Completo", value=curr_user['full_name'])
                    new_email = st.text_input("Mi Correo de Notificaciones", value=curr_user['email'] or "")
                    new_pass_val = st.text_input("Nueva Contraseña de Acceso (Opcional, dejar en blanco para no cambiar)", type="password")
                    
                    st.markdown("🔒 **Para autorizar y guardar estos cambios, ingresa la clave de seguridad:**")
                    security_key = st.text_input("Clave de Seguridad", type="password", key="sec_key_prof")
                    
                    btn_save_p = st.form_submit_button("Guardar Cambios 💾")
                    if btn_save_p:
                        if security_key != "1604":
                            st.error("❌ Clave de seguridad incorrecta. No se guardaron los cambios.")
                        elif not new_full_name or not new_email:
                            st.error("❌ El nombre y el correo electrónico son obligatorios.")
                        else:
                            conn_up_p = get_db_connection()
                            if new_pass_val.strip():
                                conn_up_p.execute("""
                                    UPDATE users 
                                    SET full_name = ?, email = ?, password = ? 
                                    WHERE username = ?
                                """, (new_full_name, new_email, new_pass_val.strip(), st.session_state.user_name))
                            else:
                                conn_up_p.execute("""
                                    UPDATE users 
                                    SET full_name = ?, email = ? 
                                    WHERE username = ?
                                """, (new_full_name, new_email, st.session_state.user_name))
                            conn_up_p.commit()
                            conn_up_p.close()
                            
                            st.session_state.full_name = new_full_name
                            # Indicar que se debe limpiar el pin en el siguiente rerun (antes de instanciar el widget)
                            st.session_state['clear_profile_pin'] = True
                            log_audit("SISTEMA", st.session_state.user_name, role, f"Actualizó datos de su perfil de usuario ({st.session_state.user_name})")
                            st.success("🎉 ¡Cambios guardados con éxito en tu perfil!")
                            st.rerun()
                            
        st.markdown("---")
        
        # Mostrar Directorio de Usuarios de forma ejecutiva como Tarjetas Nativas
        st.markdown("##### Directorio Oficial de DC Control")
        
        conn = get_db_connection()
        users_list = conn.execute("SELECT * FROM users").fetchall()
        conn.close()
        
        for u in users_list:
            # Crear contenedor limpio para cada usuario
            with st.container(border=True):
                col_u_info, col_u_action = st.columns([4, 1])
                with col_u_info:
                    st.markdown(f"#### {u['full_name']} (`{u['username']}`)")
                    st.write(f"💼 **Puesto:** {u['role']}")
                    st.write(f"📧 **Correo Electrónico:** {u['email'] or 'No registrado'}")
                with col_u_action:
                    # Botón de eliminar (Sólo Admin/Director puede usarlo, no se puede eliminar a sí mismo ni a admin principal)
                    if role == "Admin/Director" and u['username'] != st.session_state.user_name and u['username'] != 'noe.ortizadm':
                        if st.button("Quitar 🗑️", key=f"del_user_{u['username']}"):
                            conn = get_db_connection()
                            conn.execute("DELETE FROM users WHERE username = ?", (u['username'],))
                            conn.commit()
                            conn.close()
                            log_audit("SISTEMA", st.session_state.full_name, role, f"Eliminó cuenta de usuario: {u['username']}")
                            st.success(f"Usuario {u['username']} eliminado.")
                            st.rerun()
                            
        st.markdown("---")
        
        # Registrar nuevo usuario (Solo Admin)
        if role == "Admin/Director":
            with st.expander("➕ Registrar Nuevo Colaborador"):
                with st.form("Add User Form"):
                    nu_user = st.text_input("Usuario (Login)")
                    nu_pass = st.text_input("Contraseña", type="password")
                    nu_name = st.text_input("Nombre Completo")
                    nu_email = st.text_input("Correo Electrónico Individual")
                    nu_role = st.selectbox("Puesto / Rol", [
                        "Ventas",
                        "Líder Regional - Sur",
                        "Líder Regional - Norte",
                        "Analista de Costos Jefe",
                        "Analista de Costos Junior 1",
                        "Analista de Costos Junior 2",
                        "Ingeniero",
                        "Director Comercial",
                        "Director de Proyectos"
                    ])
                    
                    btn_nu = st.form_submit_button("Crear Cuenta")
                    if btn_nu:
                        if not nu_user or not nu_pass or not nu_name or not nu_email:
                            st.error("Todos los campos de registro son obligatorios.")
                        else:
                            conn = get_db_connection()
                            try:
                                conn.execute("INSERT INTO users VALUES (?, ?, ?, ?, ?)", (nu_user, nu_pass, nu_name, nu_role, nu_email))
                                conn.commit()
                                log_audit("SISTEMA", st.session_state.full_name, role, f"Creó nuevo usuario: {nu_user}")
                                st.success(f"Usuario {nu_user} creado con éxito.")
                                st.rerun()
                            except psycopg2.IntegrityError:
                                st.error("El nombre de usuario ingresado ya se encuentra registrado.")
                            finally:
                                conn.close()
                                


# ==========================================
# MÓDULO ADICIONAL: CONSOLA DE CONTROL Y RESPALDOS (Solo Admin)
# ==========================================
if "⚙️ Consola de Control" in tab_dict:
    with tab_dict["⚙️ Consola de Control"]:
        st.subheader("⚙️ Consola de Control de DC Control")
        st.write("Panel exclusivo para el Administrador para controlar respaldos, mantenimientos y canales de notificación global.")
        
        col_c_left, col_c_right = st.columns(2)
        
        with col_c_left:
            st.markdown("##### 💾 Respaldo y Mantenimiento de Datos")
            with st.container(border=True):
                st.write("Genere respaldos estructurados para evitar saturar el almacenamiento de la nube. Puede descargar un respaldo plano (.json) o un archivo compilado (.zip) que contiene los documentos físicos, resúmenes y dossiers agrupados en carpetas por proyecto.")
                
                # --- Advanced ZIP Backup compilation ---
                try:
                    zip_data = generate_structured_zip_backup()
                    col_z_dl, col_z_pc = st.columns(2)
                    with col_z_dl:
                        st.download_button(
                            label="📥 Descargar ZIP Compilado (Web)",
                            data=zip_data,
                            file_name=f"respaldo_documental_dc_control_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip",
                            mime="application/zip",
                            use_container_width=True,
                            key="btn_zip_dl"
                        )
                    with col_z_pc:
                        if st.button("💾 Guardar ZIP Compilado en PC", use_container_width=True, key="btn_zip_save_pc"):
                            save_file_directly_to_pc(zip_data, f"respaldo_documental_dc_control_{datetime.now().strftime('%Y%m%d_%H%M%S')}.zip")
                except Exception as e_zip:
                    st.error(f"Error al compilar respaldo ZIP: {e_zip}")
                    
                st.markdown("---")
                
                # --- JSON backup download ---
                try:
                    conn_bk = get_db_connection()
                    users_bk = [dict(row) for row in conn_bk.execute("SELECT * FROM users").fetchall()]
                    projects_bk = [dict(row) for row in conn_bk.execute("SELECT * FROM projects").fetchall()]
                    uploads_bk = [dict(row) for row in conn_bk.execute("SELECT * FROM uploads").fetchall()]
                    audit_bk = [dict(row) for row in conn_bk.execute("SELECT * FROM audit_log").fetchall()]
                    conn_bk.close()
                    
                    # Safe Base64 encoding of binary data for JSON compatibility
                    for row_u in uploads_bk:
                        if row_u.get('file_data') is not None:
                            try:
                                row_u['file_data'] = base64.b64encode(bytes(row_u['file_data'])).decode('utf-8')
                            except:
                                row_u['file_data'] = None
                                
                    bk_data = {
                        "backup_date": datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                        "users": users_bk,
                        "projects": projects_bk,
                        "uploads": uploads_bk,
                        "audit_log": audit_bk
                    }
                    bk_json = json.dumps(bk_data, ensure_ascii=False, indent=2)
                    
                    col_bk_dl, col_bk_sv = st.columns(2)
                    with col_bk_dl:
                        st.download_button(
                            label="📥 Generar Respaldo Plano (JSON - Web)",
                            data=bk_json,
                            file_name=f"respaldo_dc_control_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json",
                            mime="application/json",
                            use_container_width=True,
                            key="btn_json_dl"
                        )
                    with col_bk_sv:
                        if st.button("💾 Guardar Respaldo Plano en PC", use_container_width=True, key="btn_json_save_pc"):
                            save_file_directly_to_pc(bk_json.encode('utf-8'), f"respaldo_dc_control_{datetime.now().strftime('%Y%m%d_%H%M%S')}.json")
                except Exception as e_bk:
                    st.error(f"Error al generar respaldo JSON: {e_bk}")
                    
                st.markdown("---")
                
                # --- Database Restore ---
                st.markdown("##### 📤 Restaurar Copia de Seguridad JSON")
                uploaded_bk_file = st.file_uploader("Seleccione un archivo de respaldo (.json)", type=["json"], key="uploader_backup_rest_v36")
                if uploaded_bk_file:
                    confirm_rest = st.checkbox("Confirmo que deseo RESTAURAR reemplazando toda la información actual de la base de datos.", key="chk_rest_confirm")
                    if st.button("Proceder con la Restauración de Datos ⚡", type="primary", disabled=not confirm_rest, use_container_width=True, key="btn_execute_restore"):
                        try:
                            bk_content = uploaded_bk_file.read().decode("utf-8")
                            bk_parsed = json.loads(bk_content)
                            
                            required_keys = ["users", "projects", "uploads", "audit_log"]
                            if not all(k in bk_parsed for k in required_keys):
                                st.error("❌ El archivo cargado no es un archivo de respaldo válido de DC Control.")
                            else:
                                conn_rest = get_db_connection()
                                cursor_rest = conn_rest.cursor()
                                cursor_rest.execute("DROP TABLE IF EXISTS projects CASCADE")
                                cursor_rest.execute("DROP TABLE IF EXISTS audit_log CASCADE")
                                cursor_rest.execute("DROP TABLE IF EXISTS uploads CASCADE")
                                cursor_rest.execute("DROP TABLE IF EXISTS users CASCADE")
                                cursor_rest.execute("DROP TABLE IF EXISTS system_settings CASCADE")
                                conn_rest.commit()
                                conn_rest.close()
                                
                                init_db(insert_demos=False)
                                
                                conn_ins = get_db_connection()
                                cursor_ins = conn_ins.cursor()
                                
                                for u in bk_parsed["users"]:
                                    cursor_ins.execute("""
                                        INSERT INTO users (username, password, full_name, role, email)
                                        VALUES (?, ?, ?, ?, ?)
                                    """, (u['username'], u['password'], u['full_name'], u['role'], u.get('email', '')))
                                    
                                for p_bk in bk_parsed["projects"]:
                                    cursor_ins.execute("""
                                        INSERT INTO projects (
                                            id, name, client, total_amount, final_amount, state, zone, 
                                            assigned_lider, assigned_costos, assigned_ventas, priority, status, current_stage, 
                                            lose_reason, lose_percentage_gap, created_at, target_date, 
                                            step1_completed, step2_ventas_done, step2_lider_done, step2_completed, 
                                            step3_completed, step4_completed, step5_completed, step6_completed
                                        ) VALUES (?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?, ?)
                                    """, (
                                        p_bk['id'], p_bk['name'], p_bk['client'], p_bk.get('total_amount', 0.0), p_bk.get('final_amount', 0.0), p_bk['state'], p_bk['zone'],
                                        p_bk['assigned_lider'], p_bk['assigned_costos'], p_bk['assigned_ventas'], p_bk.get('priority', 'Media'), p_bk['status'], p_bk['current_stage'],
                                        p_bk.get('lose_reason', ''), p_bk.get('lose_percentage_gap', 0.0), p_bk['created_at'], p_bk['target_date'],
                                        p_bk.get('step1_completed', 0), p_bk.get('step2_ventas_done', 0), p_bk.get('step2_lider_done', 0), p_bk.get('step2_completed', 0),
                                        p_bk.get('step3_completed', 0), p_bk.get('step4_completed', 0), p_bk.get('step5_completed', 0), p_bk.get('step6_completed', 0)
                                    ))
                                    
                                for up in bk_parsed["uploads"]:
                                    file_data_bytes = None
                                    if up.get('file_data') is not None:
                                        try:
                                            file_data_bytes = base64.b64decode(up['file_data'])
                                        except:
                                            pass
                                    cursor_ins.execute("""
                                        INSERT INTO uploads (id, project_id, step_name, filename, file_path, uploaded_by, uploaded_at, file_data)
                                        VALUES (?, ?, ?, ?, ?, ?, ?, ?)
                                    """, (up['id'], up['project_id'], up['step_name'], up['filename'], up['file_path'], up['uploaded_by'], up['uploaded_at'], psycopg2.Binary(file_data_bytes) if file_data_bytes else None))
                                    
                                for log in bk_parsed["audit_log"]:
                                    cursor_ins.execute("""
                                        INSERT INTO audit_log (id, project_id, user_name, role, action, timestamp)
                                        VALUES (?, ?, ?, ?, ?, ?)
                                    """, (log['id'], log['project_id'], log['user_name'], log['role'], log['action'], log['timestamp']))
                                    
                                try:
                                    cursor_ins.execute("SELECT setval(pg_get_serial_sequence('uploads', 'id'), coalesce(max(id), 1))")
                                    cursor_ins.execute("SELECT setval(pg_get_serial_sequence('audit_log', 'id'), coalesce(max(id), 1))")
                                except Exception:
                                    pass
                                    
                                conn_ins.commit()
                                conn_ins.close()
                                st.success("🎉 ¡Copia de seguridad restaurada de forma exitosa!")
                                log_audit("SISTEMA", st.session_state.full_name, role, "Restauró base de datos desde un archivo de respaldo (.json)")
                                st.rerun()
                        except Exception as e_rest:
                            st.error(f"❌ Error al restaurar respaldo: {e_rest}")

                st.markdown("---")
                
                # --- Database wipe / reset ---
                st.markdown("##### 🚨 Restablecer Base de Datos a Cero")
                st.warning("Esta acción borrará de manera definitiva todos los proyectos, archivos, bitácoras de auditoría y reportes en Supabase. Las cuentas de usuario y contraseñas permanecerán seguras.")
                confirm_reset = st.checkbox("Entiendo los efectos y deseo limpiar a cero toda la base de datos.", key="chk_reset_confirm_v36")
                if st.button("Restablecer Base de Datos a Cero ⚠️", type="primary", disabled=not confirm_reset, key="btn_execute_wipe"):
                    try:
                        conn_res = get_db_connection()
                        cursor_res = conn_res.cursor()
                        cursor_res.execute("DROP TABLE IF EXISTS projects CASCADE")
                        cursor_res.execute("DROP TABLE IF EXISTS audit_log CASCADE")
                        cursor_res.execute("DROP TABLE IF EXISTS uploads CASCADE")
                        conn_res.commit()
                        conn_res.close()
                        
                        if os.path.exists(UPLOAD_DIR):
                            shutil.rmtree(UPLOAD_DIR)
                            os.makedirs(UPLOAD_DIR)
                            
                        init_db(insert_demos=False)
                        st.success("¡Base de datos de DC Control restablecida a cero de manera exitosa!")
                        st.rerun()
                    except Exception as e:
                        st.error(f"Error al restablecer base de datos: {e}")

        with col_c_right:
            st.markdown("##### ✉️ Canales de Notificaciones y Alertas (SMTP & Teams)")
            with st.container(border=True):
                st.write("Configure las credenciales de correo de su empresa para despachar avisos automáticos a los involucrados cada vez que se complete un paso del flujo secuencial de cotizaciones.")
                
                # Read existing configs
                notif_enabled = get_system_setting("notifications_enabled", "0")
                curr_host = get_system_setting("smtp_host", "smtp.gmail.com")
                curr_port = get_system_setting("smtp_port", "587")
                curr_user = get_system_setting("smtp_user", "notificaciones@dccontrol.com")
                curr_pass = get_system_setting("smtp_pass", "")
                curr_sender = get_system_setting("smtp_sender", "DC Control Notificaciones")
                curr_teams = get_system_setting("teams_webhook_url", "")
                
                enable_notif = st.checkbox("Habilitar Notificaciones de Sistema", value=(notif_enabled == "1"), key="cfg_notif_enabled")
                edit_host = st.text_input("Servidor SMTP (Host)", value=curr_host, key="cfg_smtp_host")
                edit_port = st.text_input("Puerto SMTP", value=curr_port, key="cfg_smtp_port")
                edit_user = st.text_input("Correo Emisor (SMTP User)", value=curr_user, key="cfg_smtp_user")
                edit_pass = st.text_input("Contraseña / Clave de Aplicación", value=curr_pass, type="password", key="cfg_smtp_pass")
                edit_sender = st.text_input("Nombre de Remitente Visible", value=curr_sender, key="cfg_smtp_sender")
                edit_teams = st.text_input("Microsoft Teams Webhook URL (Canal o Chat)", value=curr_teams, key="cfg_teams_webhook")
                
                col_cfg_save, col_cfg_test = st.columns(2)
                with col_cfg_save:
                    if st.button("Guardar Configuración 💾", use_container_width=True, type="primary", key="cfg_save_btn"):
                        set_system_setting("notifications_enabled", "1" if enable_notif else "0")
                        set_system_setting("smtp_host", edit_host)
                        set_system_setting("smtp_port", edit_port)
                        set_system_setting("smtp_user", edit_user)
                        set_system_setting("smtp_pass", edit_pass)
                        set_system_setting("smtp_sender", edit_sender)
                        set_system_setting("teams_webhook_url", edit_teams)
                        st.success("🎉 ¡Configuración guardada y sincronizada globalmente!")
                        st.rerun()
                with col_cfg_test:
                    if st.button("Enviar Correo de Prueba ✉️", use_container_width=True, key="cfg_test_btn"):
                        # Lookup admin's email
                        conn = get_db_connection()
                        admin_user = conn.execute("SELECT email FROM users WHERE username = 'noe.ortizadm'").fetchone()
                        conn.close()
                        admin_email = admin_user['email'] if admin_user else "director@dccontrol.com"
                        
                        try:
                            msg = MIMEMultipart()
                            msg['From'] = f"{edit_sender} <{edit_user}>"
                            msg['To'] = admin_email
                            msg['Subject'] = "🏗️ DC Control - Validación SMTP Exitosa"
                            body_html = f"""<html>
                            <body style="font-family: Arial, sans-serif; color: #333333;">
                                <div style="background-color: #111827; color: white; padding: 15px 20px; border-radius: 6px 6px 0 0; border-left: 6px solid #00C875;">
                                    <h2 style="margin: 0; font-size: 18px;">🏗️ Validación de Consola de Control - DC Control</h2>
                                </div>
                                <div style="padding: 20px; border: 1px solid #e5e7eb; border-top: none; border-radius: 0 0 6px 6px;">
                                    <p>¡Hola <strong>Noe Ortiz</strong>!</p>
                                    <p>Este es un correo de prueba enviado desde tu nueva <strong>Consola de Control de Escritorio</strong>.</p>
                                    <p>La configuración del servidor SMTP y el envío global de notificaciones han sido validados con éxito. El sistema ya está listo para alertar a tu equipo técnico y comercial en tiempo real.</p>
                                    <hr style="border: none; border-top: 1px solid #e5e7eb; margin: 15px 0;">
                                    <p style="font-size: 11px; color: #6b7280; text-align: center;">DC Control S.A. de C.V. • Trazabilidad y Gobernabilidad</p>
                                </div>
                            </body>
                            </html>"""
                            msg.attach(MIMEText(body_html, 'html'))
                            
                            server = smtplib.SMTP(edit_host, int(edit_port))
                            server.starttls()
                            server.login(edit_user, edit_pass)
                            server.sendmail(edit_user, admin_email, msg.as_string())
                            server.quit()
                            st.success(f"📬 ¡Correo de prueba enviado con éxito a {admin_email}! Revisa tu bandeja de entrada.")
                        except Exception as e_test_m:
                            st.error(f"❌ Error al enviar correo de prueba: {e_test_m}")

# ==========================================
# MÓDULO 6: BITÁCORA DE AUDITORÍA (Solo Admin)
# ==========================================
if "📜 Bitácora Auditoría" in tab_dict:
    with tab_dict["📜 Bitácora Auditoría"]:
        st.subheader("📜 Bitácora de Trazabilidad")
        st.write("Registro histórico de acciones realizadas en la base de datos:")
        
        conn = get_db_connection()
        logs = conn.execute("SELECT * FROM audit_log ORDER BY timestamp DESC").fetchall()
        conn.close()
        
        if not logs:
            st.info("No se han registrado acciones comerciales ni técnicas de momento.")
        else:
            df_logs = pd.DataFrame([dict(l) for l in logs])
            df_logs_display = df_logs.rename(columns={
                'project_id': 'ID Proyecto',
                'user_name': 'Usuario Responsable',
                'role': 'Puesto / Rol',
                'action': 'Acción Realizada',
                'timestamp': 'Fecha y Hora'
            })[['ID Proyecto', 'Usuario Responsable', 'Puesto / Rol', 'Acción Realizada', 'Fecha y Hora']]
            
            # Obtener mapeo de IDs de proyecto a nombres para mostrar en el desplegable
            conn_map = get_db_connection()
            all_projs = conn_map.execute("SELECT id, name FROM projects").fetchall()
            conn_map.close()
            
            p_id_to_name = {p['id']: f"{p['id']} - {p['name']}" for p in all_projs}
            p_id_to_name['SISTEMA'] = 'SISTEMA'
            p_id_to_name[''] = 'Sin ID'
            
            # Construir opciones únicas presentes en los logs
            unique_ids = df_logs['project_id'].unique()
            log_filter_options = ["Todos"] + sorted([p_id_to_name.get(pid, pid) for pid in unique_ids if pid])
            
            selected_log_filter = st.selectbox(
                "🔍 Filtrar Bitácora por Obra / ID de Proyecto:",
                log_filter_options,
                key="bitacora_project_filter"
            )
            
            if selected_log_filter != "Todos":
                filter_id = selected_log_filter.split(" - ")[0]
                df_logs_display = df_logs_display[df_logs_display['ID Proyecto'] == filter_id]
                
            st.dataframe(df_logs_display, use_container_width=True, hide_index=True)

# Footer Corporativo
st.markdown("---")
st.markdown(
    "<p style='text-align: center; color: #9ca3af; font-size: 11px;'> "
    "DC Control S.A. de C.V. • Control de Cotizaciones"
    "</p>", 
    unsafe_allow_html=True
)
